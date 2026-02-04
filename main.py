import subprocess
from flask import Flask, request, send_file, jsonify
from docxtpl import DocxTemplate
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import tempfile
import fitz  # PyMuPDF
from PIL import Image
import io
import json
import traceback
from flask_cors import CORS

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})
# --- ฟังก์ชันแปลง docx → pdf ด้วย LibreOffice ---
def convert_docx_to_pdf(docx_path, output_pdf_path):
    cmd = [
        "libreoffice",
        "--headless",
        "--convert-to", "pdf",
        "--outdir", os.path.dirname(output_pdf_path),
        docx_path
    ]
    subprocess.run(cmd, check=True)

# --- ฟังก์ชันแปลงตัวเลขเป็นเลขไทย ---
def to_thai_digits(text):
    thai_digits = '๐๑๒๓๔๕๖๗๘๙'
    def convert_char(c):
        if '0' <= c <= '9':
            return thai_digits[ord(c) - ord('0')]
        return c
    if not isinstance(text, str):
        return text
    return ''.join(convert_char(c) for c in text)

# --- ฟังก์ชันวาดข้อความเป็นภาพ ---
def draw_text_image(text, font_path, font_size=20, color=(2, 53, 139), scale=1):
    from PIL import ImageFont, ImageDraw
    big_font_size = font_size * scale
    font = ImageFont.truetype(font_path, big_font_size)
    padding = 4 * scale
    lines = text.split('\n')
    width = max([font.getbbox(line)[2] for line in lines]) + 2 * padding
    height = sum([font.getbbox(line)[3] - font.getbbox(line)[1] for line in lines]) + 2 * padding + (len(lines)-1)*2*scale
    img = Image.new("RGBA", (width, height), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)
    y = padding
    for line in lines:
        draw.text((padding, y), line, font=font, fill=color)
        y += font.getbbox(line)[3] - font.getbbox(line)[1] + 2*scale
    return img


# --- สร้าง PDF จาก template docx ---
@app.route('/pdf', methods=['POST'])
def generate_pdf():
    try:
        data = request.json or {}

        # ===== เพิ่มส่วนนี้ =====
        required_fields = [
            "date",
            "subject",
            "introduction",
            "author_name",
            "author_position",
            "fact",
            "proposal"
        ]
        # ตรวจสอบฟิลด์ที่ขาด หรือ ฟิลด์ที่เป็น "" (ว่าง)
        missing = [f for f in required_fields if not data.get(f)]
        if missing:
            return jsonify({'error': f"Missing fields: {', '.join(missing)}"}), 400
        # =====================

        # ลบช่องว่างท้ายบรรทัดใน introduction, fact, proposal
        for field in ['introduction', 'fact', 'proposal']:
            if field in data and data[field]:
                # ลบช่องว่างท้ายแต่ละบรรทัด
                lines = data[field].split('\n')
                data[field] = '\n'.join(line.rstrip() for line in lines)

        # จัดรูปแบบ proposal:
        # ! = ขึ้นบรรทัดใหม่ + indent + "- "
        # ? = ขึ้นบรรทัดใหม่ + บรรทัดก่อนหน้าไม่ justify (ใช้ marker \u200B)
        if 'proposal' in data and data['proposal']:
            proposal_text = data['proposal']
            lines = []
            current_line = ""
            mark_previous_no_justify = False  # flag สำหรับ mark บรรทัดก่อนหน้า
            i = 0
            while i < len(proposal_text):
                if proposal_text[i] == '!' and i > 0:
                    # เจอ ! = ขึ้นบรรทัดใหม่ + indent + "- "
                    if current_line.strip():
                        line_text = current_line.rstrip()
                        if mark_previous_no_justify:
                            line_text = '\u200B' + line_text
                            mark_previous_no_justify = False
                        lines.append(line_text)
                    current_line = "- "
                    i += 1
                    while i < len(proposal_text) and proposal_text[i] == ' ':
                        i += 1
                    continue
                elif proposal_text[i] == '?' and i > 0:
                    # เจอ ? = ขึ้นบรรทัดใหม่ + บรรทัดนี้ไม่ justify
                    if current_line.strip():
                        # เติม marker ที่บรรทัดนี้ (บรรทัดก่อน ?)
                        lines.append('\u200B' + current_line.rstrip())
                    current_line = ""
                    i += 1
                    while i < len(proposal_text) and proposal_text[i] == ' ':
                        i += 1
                    continue
                else:
                    current_line += proposal_text[i]
                    i += 1

            if current_line.strip():
                line_text = current_line.rstrip()
                if mark_previous_no_justify:
                    line_text = '\u200B' + line_text
                lines.append(line_text)

            # รวมผลลัพธ์เป็น list สำหรับ template loop
            if lines:
                if lines[0].startswith('! '):
                    lines[0] = '- ' + lines[0][2:]
                elif lines[0].startswith('!'):
                    lines[0] = '- ' + lines[0][1:]
                data['proposal_lines'] = lines
            else:
                data['proposal_lines'] = [proposal_text]
        else:
            data['proposal_lines'] = []

        # แปลงเลขใน dict เป็นเลขไทย
        def convert_dict(d):
            if isinstance(d, dict):
                return {k: convert_dict(v) for k, v in d.items()}
            elif isinstance(d, list):
                return [convert_dict(i) for i in d]
            elif isinstance(d, str):
                return to_thai_digits(d)
            else:
                return d
        data = convert_dict(data)
        template_path = os.path.join(os.path.dirname(__file__), "templates", "memo-template2.docx")
        if not os.path.exists(template_path):
            return jsonify({'error': f'Template file not found: {template_path}'}), 500
        doc = DocxTemplate(template_path)
        doc.render(data)

        # บังคับ justify ทุก paragraph (รวมใน table ด้วย)
        # ยกเว้น paragraph ที่มี marker \u200B (ไม่ justify)
        for paragraph in doc.paragraphs:
            if '\u200B' in paragraph.text:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.text = run.text.replace('\u200B', '')
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if '\u200B' in paragraph.text:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            for run in paragraph.runs:
                                run.text = run.text.replace('\u200B', '')
                        else:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
            doc.save(tmp_docx.name)
            tmp_pdf = tmp_docx.name.replace('.docx', '.pdf')
            convert_docx_to_pdf(tmp_docx.name, tmp_pdf)

        # เพิ่มหน้าเปล่า 1 หน้าสำหรับพื้นที่ลายเซ็น
        pdf = fitz.open(tmp_pdf)
        pdf.new_page(width=pdf[0].rect.width, height=pdf[0].rect.height)
        tmp_pdf_with_blank = tmp_pdf.replace('.pdf', '_blank.pdf')
        pdf.save(tmp_pdf_with_blank)
        pdf.close()

        return send_file(tmp_pdf_with_blank, mimetype="application/pdf", as_attachment=True, download_name="memo.pdf")
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

# --- วางลายเซ็น/ความเห็นลง PDF ที่อัพโหลดมา ---
@app.route('/add_signature', methods=['POST'])
def add_signature():
    try:
        DEFAULT_SIGNATURE_HEIGHT = 70
        DEFAULT_COMMENT_FONT_SIZE = 20
        font_path = os.path.join(os.path.dirname(__file__), "fonts", "THSarabunNew.ttf")
        if not os.path.isfile(font_path):
            return jsonify({'error': f"Font file not found: {font_path}"}), 500

        if 'pdf' not in request.files:
            return jsonify({'error': 'No PDF file uploaded'}), 400
        pdf_file = request.files['pdf']

        if 'signatures' not in request.form:
            return jsonify({'error': 'No signatures data'}), 400
        signatures = json.loads(request.form['signatures'])

        pdf_bytes = pdf_file.read()
        pdf = fitz.open(stream=pdf_bytes, filetype="pdf")

        # --- กลุ่ม sig ตามตำแหน่ง (page, x, y) ---
        from collections import defaultdict
        sig_dict = defaultdict(list)
        for sig in signatures:
            page_number = int(sig.get('page', 0))
            x = int(sig['x'])
            y = int(sig['y'])
            sig_dict[(page_number, x, y)].append(sig)

        # --- วาดลายเซ็นและความเห็นทีละจุด ---
        for (page_number, x, y), sigs in sig_dict.items():
            page = pdf[page_number]
            sigs_sorted = sorted(sigs, key=lambda s: 0 if s['type'] == 'text' else 1)
            current_y = y
            for sig in sigs_sorted:
                if sig['type'] == 'text':
                    text = to_thai_digits(sig.get('text', ''))
                    font_size = DEFAULT_COMMENT_FONT_SIZE
                    orig_color = sig.get('color', (2, 53, 139))
                    if isinstance(orig_color, (list, tuple)):
                        r = min(int(orig_color[0]*0.8), 255)
                        g = min(int(orig_color[1]*0.8), 255)
                        b = min(int(orig_color[2]*0.8), 255)
                        color = (r, g, b)
                    else:
                        color = (2, 53, 139)
                    img = draw_text_image(text, font_path, font_size=font_size, color=color, scale=1)
                    img_byte_arr = io.BytesIO()
                    img.save(img_byte_arr, format='PNG')
                    rect = fitz.Rect(x, current_y, x + img.width, current_y + img.height)
                    page.insert_image(rect, stream=img_byte_arr.getvalue())
                    current_y += img.height
                elif sig['type'] == 'image':
                    file_key = sig['file_key']
                    if file_key not in request.files:
                        continue
                    signature_file = request.files[file_key]
                    img = Image.open(signature_file)
                    fixed_height = DEFAULT_SIGNATURE_HEIGHT
                    ratio = fixed_height / img.height
                    new_width = int(img.width * ratio)
                    img = img.resize((new_width, fixed_height), resample=Image.LANCZOS)
                    img_byte_arr = io.BytesIO()
                    img.save(img_byte_arr, format='PNG')
                    rect = fitz.Rect(x, current_y, x + new_width, current_y + fixed_height)
                    page.insert_image(rect, stream=img_byte_arr.getvalue())
                    current_y += fixed_height

        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_pdf:
            pdf.save(tmp_pdf.name)
        pdf.close()
        return send_file(tmp_pdf.name, mimetype="application/pdf", as_attachment=True, download_name="signed.pdf")
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({'error': str(e)}), 500


@app.route('/add_signature_v2', methods=['POST'])
def add_signature_v2():
    # --- ฟังก์ชันวาดข้อความเป็นภาพ (v2) ---
    def draw_text_image_v2(text, font_path, font_size=20, color=(2, 53, 139), scale=1, font_weight="regular", line_height_ratio=1.2):
        from PIL import ImageFont, ImageDraw, Image
        # เลือก font ตาม font_weight
        if font_weight == "bold":
            font_path = os.path.join(os.path.dirname(__file__), "fonts", "THSarabunNew Bold.ttf")
        big_font_size = font_size * scale
        font = ImageFont.truetype(font_path, big_font_size)
        padding = 4 * scale
        lines = text.split('\n')

        # ใช้ fixed line height แทน bbox เพื่อหลีกเลี่ยงปัญหา tone marks ทำให้ความสูงไม่เท่ากัน
        # รองรับการปรับ line_height_ratio (default 1.2, สำหรับ comment ใช้ 0.96 = 1.2 * 0.8)
        fixed_line_height = int(font_size * line_height_ratio * scale)

        # วัดความกว้างเท่านั้น
        dummy_img = Image.new("RGBA", (10, 10), (255, 255, 255, 0))
        dummy_draw = ImageDraw.Draw(dummy_img)
        line_widths = []
        for line in lines:
            bbox = dummy_draw.textbbox((0, 0), line, font=font)
            width = bbox[2] - bbox[0]
            line_widths.append(width)

        max_width = max(line_widths) + 2 * padding
        total_height = len(lines) * fixed_line_height + 2 * padding

        img = Image.new("RGBA", (max_width, total_height), (255, 255, 255, 0))
        draw = ImageDraw.Draw(img)

        y = padding
        for line in lines:
            # ใช้ anchor="la" (left-ascender) เพื่อยึดตำแหน่งที่ ascender line
            # ทำให้ทุกบรรทัดวางที่ตำแหน่งเดียวกัน ไม่ว่าจะมี tone marks หรือไม่
            draw.text((padding, y), line, font=font, fill=color, anchor="la")
            y += fixed_line_height

        return img
    try:
        DEFAULT_SIGNATURE_HEIGHT = 50
        DEFAULT_COMMENT_FONT_SIZE = 18
        LINE_SPACING = 20  # ระยะห่างคงที่ระหว่างบรรทัดข้อความ (สำหรับ single-line text)
        font_path = os.path.join(os.path.dirname(__file__), "fonts", "THSarabunNew.ttf")
        if not os.path.isfile(font_path):
            return jsonify({'error': f"Font file not found: {font_path}"}), 500

        if 'pdf' not in request.files:
            return jsonify({'error': 'No PDF file uploaded'}), 400
        pdf_file = request.files['pdf']

        if 'signatures' not in request.form:
            return jsonify({'error': 'No signatures data'}), 400
        signatures = json.loads(request.form['signatures'])

        pdf_bytes = pdf_file.read()
        pdf = fitz.open(stream=pdf_bytes, filetype="pdf")

        from collections import defaultdict
        sig_dict = defaultdict(list)
        for sig in signatures:
            page_number = int(sig.get('page', 0))
            x = int(sig['x'])
            y = int(sig['y'])
            # รองรับ width/height สำหรับ center positioning
            width = sig.get('width', 0)
            height = sig.get('height', 0)
            
            # ถ้าไม่มี width/height ให้ใช้ค่า default สำหรับ center positioning
            if width == 0 and height == 0:
                width = 120  # default width
                height = 60  # default height
                print(f"DEBUG: Using default dimensions {width}x{height} for signature at ({x}, {y})")
            
            sig_dict[(page_number, x, y, width, height)].append(sig)

        for (page_number, x, y, width, height), sigs in sig_dict.items():
            page = pdf[page_number]
            
            # Debug: แสดงข้อมูล page และพิกัด
            page_rect = page.rect
            print(f"DEBUG: Page {page_number} - Size: {page_rect.width}x{page_rect.height}")
            print(f"DEBUG: Original coordinates: ({x}, {y})")
            print(f"DEBUG: Signature dimensions: {width}x{height}")
            print(f"DEBUG: Page bounds: x(0-{page_rect.width}), y(0-{page_rect.height})")
            
            # ถ้ามี width/height แสดงว่าเป็น center positioning
            is_center_positioning = width > 0 and height > 0
            
            # Logic สำหรับปรับพิกัด Y ให้กลับกัน (บนเป็นล่าง ล่างเป็นบน)
            # คำนวณ: new_y = (page_height - original_y - signature_height) + height_offset
            if is_center_positioning:
                # สำหรับ center positioning ใช้ height ที่กำหนดมา
                adjusted_y = page_rect.height - y - height
                # เลื่อนลงแนวดิ่งเท่ากับ height (60)
                adjusted_y += height+30
                print(f"DEBUG: Y-axis flip with center positioning: {y} -> {adjusted_y} (with +{height} offset)")
                center_x = x
                center_y = adjusted_y
                print(f"DEBUG: Using center positioning - adjusted coordinates")
                print(f"DEBUG: Center point: ({center_x}, {center_y})")
                print(f"DEBUG: Bounding box: {width}x{height}")
            else:
                # สำหรับ top-left positioning ใช้ default signature height
                signature_box_height = 60  # default height สำหรับการคำนวณ
                adjusted_y = page_rect.height - y - signature_box_height
                # เลื่อนลงแนวดิ่งเท่ากับ 60
                adjusted_y += 60
                print(f"DEBUG: Y-axis flip with top-left positioning: {y} -> {adjusted_y} (with +60 offset)")
                center_x = x
                center_y = adjusted_y
                print(f"DEBUG: Using top-left positioning - adjusted coordinates")
            
            current_y = center_y  # ใช้ค่า Y ที่ปรับแล้ว
            # Check if any signature has 'lines' field
            has_lines = any('lines' in sig for sig in sigs)
            if has_lines:
                # For each sig with lines, draw lines in order
                for sig in sigs:
                    lines = sig.get('lines')
                    if not lines:
                        # fallback to old logic for this sig
                        if sig['type'] == 'text':
                            text = to_thai_digits(sig.get('text', ''))
                            # ถ้า type == "comment" ให้ font_size=20, font_weight="bold", line_height_ratio=0.96
                            font_size = 20 if sig.get('type') == 'comment' else DEFAULT_COMMENT_FONT_SIZE
                            font_weight = "bold" if sig.get('type') == 'comment' else "regular"
                            line_height_ratio = 0.96 if sig.get('type') == 'comment' else 1.2
                            orig_color = sig.get('color', (2, 53, 139))
                            if isinstance(orig_color, (list, tuple)):
                                r = min(int(orig_color[0]*0.8), 255)
                                g = min(int(orig_color[1]*0.8), 255)
                                b = min(int(orig_color[2]*0.8), 255)
                                color = (r, g, b)
                            else:
                                color = (2, 53, 139)
                            img = draw_text_image_v2(text, font_path, font_size=font_size, color=color, scale=1, font_weight=font_weight, line_height_ratio=line_height_ratio)
                            img_byte_arr = io.BytesIO()
                            img.save(img_byte_arr, format='PNG')
                            # ใช้ center positioning ถ้ามี width/height
                            if is_center_positioning:
                                left_x = center_x - img.width // 2
                                top_y = center_y - img.height // 2
                            else:
                                left_x = x
                                top_y = current_y
                            rect = fitz.Rect(left_x, top_y, left_x + img.width, top_y + img.height)
                            print(f"DEBUG: Text '{text}' placed at rect: {rect} (center_pos: {is_center_positioning})")
                            page.insert_image(rect, stream=img_byte_arr.getvalue())
                            if not is_center_positioning:
                                num_lines = text.count('\n') + 1
                                if num_lines == 1:
                                    current_y += LINE_SPACING
                                else:
                                    current_y += img.height + 4
                        elif sig['type'] == 'image':
                            file_key = sig['file_key']
                            if file_key not in request.files:
                                continue
                            signature_file = request.files[file_key]
                            img = Image.open(signature_file)
                            fixed_height = DEFAULT_SIGNATURE_HEIGHT
                            ratio = fixed_height / img.height
                            new_width = int(img.width * ratio)
                            img = img.resize((new_width, fixed_height), resample=Image.LANCZOS)
                            img_byte_arr = io.BytesIO()
                            img.save(img_byte_arr, format='PNG')
                            # ใช้ center positioning ถ้ามี width/height
                            if is_center_positioning:
                                left_x = center_x - new_width // 2
                                top_y = center_y - fixed_height // 2
                            else:
                                left_x = x
                                top_y = current_y
                            rect = fitz.Rect(left_x, top_y, left_x + new_width, top_y + fixed_height)
                            print(f"DEBUG: Image placed at rect: {rect} (center_pos: {is_center_positioning})")
                            page.insert_image(rect, stream=img_byte_arr.getvalue())
                            if not is_center_positioning:
                                current_y += fixed_height - 10  # ลดระยะห่างให้ใกล้กับข้อความด้านล่าง
                    else:
                        # draw lines in order - รองรับ center positioning
                        if is_center_positioning:
                            # สำหรับ center positioning ให้เริ่มจากด้านบนของ bounding box
                            current_y = center_y - height // 2
                        
                        for line in lines:
                            line_type = line.get('type')
                            if line_type == 'image':
                                file_key = line.get('file_key')
                                if file_key and file_key in request.files:
                                    signature_file = request.files[file_key]
                                    img = Image.open(signature_file)
                                    fixed_height = DEFAULT_SIGNATURE_HEIGHT
                                    ratio = fixed_height / img.height
                                    new_width = int(img.width * ratio)
                                    img = img.resize((new_width, fixed_height), resample=Image.LANCZOS)
                                    img_byte_arr = io.BytesIO()
                                    img.save(img_byte_arr, format='PNG')
                                    
                                    if is_center_positioning:
                                        left_x = center_x - new_width // 2  # คืนค่าเดิม แต่เพิ่ม debug
                                        top_y = current_y
                                        print(f"DEBUG: Image center positioning - center_x:{center_x}, new_width:{new_width}, left_x:{left_x}")
                                        print(f"DEBUG: Expected position - should place image at left edge: {left_x}")
                                    else:
                                        left_x = x
                                        top_y = current_y

                                    rect = fitz.Rect(left_x, top_y, left_x + new_width, top_y + fixed_height)
                                    print(f"DEBUG: Image rect: {rect}")
                                    page.insert_image(rect, stream=img_byte_arr.getvalue())
                                    current_y += fixed_height - 10  # ลดระยะห่างให้ใกล้กับข้อความด้านล่าง
                            else:
                                # For text types: 'comment', 'name', 'position', 'academic_rank', 'org_structure_role', 'timestamp'
                                text_value = line.get('text') or line.get('value') or line.get('comment') or ''

                                # ฟังก์ชันนับตัวอักษรที่มองเห็น (ไม่รวม tone marks, vowel marks)
                                def count_visible_chars(s):
                                    thai_marks = set([
                                        '\u0E31', '\u0E34', '\u0E35', '\u0E36', '\u0E37',
                                        '\u0E38', '\u0E39', '\u0E3A', '\u0E47', '\u0E48',
                                        '\u0E49', '\u0E4A', '\u0E4B', '\u0E4C', '\u0E4D', '\u0E4E'
                                    ])
                                    return len([c for c in s if c not in thai_marks])

                                # ฟังก์ชันตัดข้อความตามจำนวนตัวอักษรที่มองเห็น
                                def wrap_by_visible_chars(text, max_chars=30):
                                    if count_visible_chars(text) <= max_chars:
                                        return [text]

                                    lines = []
                                    current = ""
                                    for char in text:
                                        test = current + char
                                        if count_visible_chars(test) <= max_chars:
                                            current = test
                                        else:
                                            if current:
                                                lines.append(current)
                                            current = char
                                    if current:
                                        lines.append(current)
                                    return lines

                                # ถ้าเป็น comment และมี - ให้แยกเป็นหลายบรรทัด
                                if line_type == 'comment' and '-' in text_value:
                                    # แยกตาม - และเก็บ - ไว้ข้างหน้าแต่ละบรรทัด
                                    parts = text_value.split('-')
                                    # เอา parts ที่ไม่ว่างเปล่าออกมา
                                    text_lines = ['-' + part for part in parts if part.strip()]

                                    # ตัดแต่ละบรรทัดให้ไม่เกิน 30 ตัวอักษรที่มองเห็น
                                    wrapped_lines = []
                                    for tline in text_lines:
                                        wrapped = wrap_by_visible_chars(tline, max_chars=30)
                                        wrapped_lines.extend(wrapped)

                                    text = '\n'.join([to_thai_digits(t) for t in wrapped_lines])
                                else:
                                    text = to_thai_digits(text_value)
                                # ถ้า type == "comment" ให้ font_size=20, font_weight="bold", line_height_ratio=0.96 (ลด 20%)
                                font_size = 20 if line_type == 'comment' else DEFAULT_COMMENT_FONT_SIZE
                                font_weight = "bold" if line_type == 'comment' else "regular"
                                line_height_ratio = 0.96 if line_type == 'comment' else 1.2  # comment ใช้ 0.96 = 1.2 * 0.8
                                orig_color = line.get('color', (2, 53, 139))
                                if isinstance(orig_color, (list, tuple)):
                                    r = min(int(orig_color[0]*0.8), 255)
                                    g = min(int(orig_color[1]*0.8), 255)
                                    b = min(int(orig_color[2]*0.8), 255)
                                    color = (r, g, b)
                                else:
                                    color = (2, 53, 139)
                                img = draw_text_image_v2(text, font_path, font_size=font_size, color=color, scale=1, font_weight=font_weight, line_height_ratio=line_height_ratio)
                                img_byte_arr = io.BytesIO()
                                img.save(img_byte_arr, format='PNG')
                                
                                if is_center_positioning:
                                    left_x = center_x - img.width // 2  # คืนค่าเดิม
                                    top_y = current_y
                                    print(f"DEBUG: Text center positioning - center_x:{center_x}, img.width:{img.width}, left_x:{left_x}")
                                    print(f"DEBUG: Expected position - should place text at left edge: {left_x}")
                                else:
                                    left_x = x
                                    top_y = current_y

                                rect = fitz.Rect(left_x, top_y, left_x + img.width, top_y + img.height)
                                print(f"DEBUG: Text '{text}' rect: {rect}")
                                num_lines = text.count('\n') + 1
                                print(f"DEBUG: Text has {num_lines} lines, img.height={img.height}")
                                page.insert_image(rect, stream=img_byte_arr.getvalue())
                                # ถ้าเป็นบรรทัดเดียวใช้ fixed spacing, ถ้าหลายบรรทัดใช้ความสูงจริง
                                if num_lines == 1:
                                    current_y += LINE_SPACING
                                else:
                                    current_y += img.height + 4
            else:
                # fallback to old logic
                sigs_sorted = sorted(sigs, key=lambda s: 0 if s['type'] == 'text' else 1)
                for sig in sigs_sorted:
                    if sig['type'] == 'text':
                        text = to_thai_digits(sig.get('text', ''))
                        # ถ้า type == "comment" ให้ font_size=20, font_weight="bold", line_height_ratio=0.96
                        font_size = 20 if sig.get('type') == 'comment' else DEFAULT_COMMENT_FONT_SIZE
                        font_weight = "bold" if sig.get('type') == 'comment' else "regular"
                        line_height_ratio = 0.96 if sig.get('type') == 'comment' else 1.2
                        orig_color = sig.get('color', (2, 53, 139))
                        if isinstance(orig_color, (list, tuple)):
                            r = min(int(orig_color[0]*0.8), 255)
                            g = min(int(orig_color[1]*0.8), 255)
                            b = min(int(orig_color[2]*0.8), 255)
                            color = (r, g, b)
                        else:
                            color = (2, 53, 139)
                        img = draw_text_image_v2(text, font_path, font_size=font_size, color=color, scale=1, font_weight=font_weight, line_height_ratio=line_height_ratio)
                        img_byte_arr = io.BytesIO()
                        img.save(img_byte_arr, format='PNG')
                        rect = fitz.Rect(x, current_y, x + img.width, current_y + img.height)
                        page.insert_image(rect, stream=img_byte_arr.getvalue())
                        num_lines = text.count('\n') + 1
                        if num_lines == 1:
                            current_y += LINE_SPACING
                        else:
                            current_y += img.height + 4
                    elif sig['type'] == 'image':
                        file_key = sig['file_key']
                        if file_key not in request.files:
                            continue
                        signature_file = request.files[file_key]
                        img = Image.open(signature_file)
                        fixed_height = DEFAULT_SIGNATURE_HEIGHT
                        ratio = fixed_height / img.height
                        new_width = int(img.width * ratio)
                        img = img.resize((new_width, fixed_height), resample=Image.LANCZOS)
                        img_byte_arr = io.BytesIO()
                        img.save(img_byte_arr, format='PNG')
                        rect = fitz.Rect(x, current_y, x + new_width, current_y + fixed_height)
                        page.insert_image(rect, stream=img_byte_arr.getvalue())
                        current_y += fixed_height - 10  # ลดระยะห่างให้ใกล้กับข้อความด้านล่าง

        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_pdf:
            pdf.save(tmp_pdf.name)
        pdf.close()
        return send_file(tmp_pdf.name, mimetype="application/pdf", as_attachment=True, download_name="signed.pdf")
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({'error': str(e)}), 500



@app.route('/2in1memo', methods=['POST'])
def generate_2in1_memo():
    """รวมการทำงานของ /pdf และ /add_signature_v2 ในครั้งเดียว"""
    try:
        # ตรวจสอบข้อมูลที่ส่งมา
        if not request.form and not request.json:
            return jsonify({'error': 'No data provided'}), 400
        
        # ข้อมูลสำหรับสร้าง PDF จาก form หรือ json
        if request.form:
            # กรณีส่งมาแบบ multipart/form-data
            data = {}
            for key in request.form:
                if key != 'signatures':  # signatures จะจัดการแยก
                    data[key] = request.form[key]
        else:
            # กรณีส่งมาแบบ JSON
            data = request.json or {}

        # ===== ส่วนที่ 1: สร้าง PDF (จาก /pdf) =====
        required_fields = [
            "doc_number",
            "date",
            "subject",
            "introduction", 
            "author_name",
            "author_position",
            "fact",
            "proposal"
        ]
        
        missing = [f for f in required_fields if not data.get(f)]
        if missing:
            return jsonify({'error': f"Missing fields: {', '.join(missing)}"}), 400

        # ลบช่องว่างท้ายบรรทัดใน introduction, fact, proposal
        for field in ['introduction', 'fact', 'proposal']:
            if field in data and data[field]:
                # ลบช่องว่างท้ายแต่ละบรรทัด
                lines = data[field].split('\n')
                data[field] = '\n'.join(line.rstrip() for line in lines)

        # จัดรูปแบบ proposal:
        # ! = ขึ้นบรรทัดใหม่ + indent + "- "
        # ? = ขึ้นบรรทัดใหม่ + บรรทัดก่อนหน้าไม่ justify (ใช้ marker \u200B)
        if 'proposal' in data and data['proposal']:
            proposal_text = data['proposal']
            lines = []
            current_line = ""
            mark_previous_no_justify = False
            i = 0
            while i < len(proposal_text):
                if proposal_text[i] == '!' and i > 0:
                    # เจอ ! = ขึ้นบรรทัดใหม่ + indent + "- "
                    if current_line.strip():
                        line_text = current_line.rstrip()
                        if mark_previous_no_justify:
                            line_text = '\u200B' + line_text
                            mark_previous_no_justify = False
                        lines.append(line_text)
                    current_line = "- "
                    i += 1
                    while i < len(proposal_text) and proposal_text[i] == ' ':
                        i += 1
                    continue
                elif proposal_text[i] == '?' and i > 0:
                    # เจอ ? = ขึ้นบรรทัดใหม่ + บรรทัดนี้ไม่ justify
                    if current_line.strip():
                        lines.append('\u200B' + current_line.rstrip())
                    current_line = ""
                    i += 1
                    while i < len(proposal_text) and proposal_text[i] == ' ':
                        i += 1
                    continue
                else:
                    current_line += proposal_text[i]
                    i += 1

            if current_line.strip():
                line_text = current_line.rstrip()
                if mark_previous_no_justify:
                    line_text = '\u200B' + line_text
                lines.append(line_text)

            # รวมผลลัพธ์เป็น list สำหรับ template loop
            if lines:
                if lines[0].startswith('! '):
                    lines[0] = '- ' + lines[0][2:]
                elif lines[0].startswith('!'):
                    lines[0] = '- ' + lines[0][1:]
                data['proposal_lines'] = lines
            else:
                data['proposal_lines'] = [proposal_text]
        else:
            data['proposal_lines'] = []

        # แปลงเลขใน dict เป็นเลขไทย
        def convert_dict(d):
            if isinstance(d, dict):
                return {k: convert_dict(v) for k, v in d.items()}
            elif isinstance(d, list):
                return [convert_dict(i) for i in d]
            elif isinstance(d, str):
                return to_thai_digits(d)
            else:
                return d
        data = convert_dict(data)
        
        template_path = os.path.join(os.path.dirname(__file__), "templates", "memo-template2.docx")
        if not os.path.exists(template_path):
            return jsonify({'error': f'Template file not found: {template_path}'}), 500
        
        doc = DocxTemplate(template_path)
        doc.render(data)

        # บังคับ justify ทุก paragraph (รวมใน table ด้วย)
        # ยกเว้น paragraph ที่มี marker \u200B (ไม่ justify)
        for paragraph in doc.paragraphs:
            if '\u200B' in paragraph.text:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.text = run.text.replace('\u200B', '')
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if '\u200B' in paragraph.text:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            for run in paragraph.runs:
                                run.text = run.text.replace('\u200B', '')
                        else:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
            doc.save(tmp_docx.name)
            tmp_pdf = tmp_docx.name.replace('.docx', '.pdf')
            convert_docx_to_pdf(tmp_docx.name, tmp_pdf)

        # เพิ่มหน้าเปล่า 1 หน้าสำหรับพื้นที่ลายเซ็น
        pdf_for_blank = fitz.open(tmp_pdf)
        pdf_for_blank.new_page(width=pdf_for_blank[0].rect.width, height=pdf_for_blank[0].rect.height)
        tmp_pdf_with_blank = tmp_pdf.replace('.pdf', '_blank.pdf')
        pdf_for_blank.save(tmp_pdf_with_blank)
        pdf_for_blank.close()
        tmp_pdf = tmp_pdf_with_blank  # ใช้ไฟล์ที่มีหน้าเปล่าต่อ

        # ===== ส่วนที่ 2: เพิ่มลายเซ็น (จาก /add_signature_v2) =====

        # ตรวจสอบว่ามี signatures หรือไม่
        if 'signatures' not in request.form:
            # ถ้าไม่มี signatures ให้ return PDF ธรรมดา
            return send_file(tmp_pdf, mimetype="application/pdf", as_attachment=True, download_name="memo.pdf")
        
        signatures = json.loads(request.form['signatures'])
        
        # ตรวจสอบว่ามีไฟล์เอกสารแนบที่ต้องการแปะลายเซ็นหรือไม่
        attachment_pdf = None
        if 'attachment_pdf' in request.files:
            attachment_file = request.files['attachment_pdf']
            attachment_pdf_bytes = attachment_file.read()
            attachment_pdf = fitz.open(stream=attachment_pdf_bytes, filetype="pdf")
        
        # ฟังก์ชันวาดข้อความเป็นภาพ (v2)
        def draw_text_image_v2(text, font_path, font_size=20, color=(2, 53, 139), scale=1, font_weight="regular"):
            from PIL import ImageFont, ImageDraw, Image
            if font_weight == "bold":
                font_path = os.path.join(os.path.dirname(__file__), "fonts", "THSarabunNew Bold.ttf")
            big_font_size = font_size * scale
            font = ImageFont.truetype(font_path, big_font_size)
            padding = 4 * scale
            lines = text.split('\n')
            dummy_img = Image.new("RGBA", (10, 10), (255, 255, 255, 0))
            dummy_draw = ImageDraw.Draw(dummy_img)
            line_sizes = []
            for line in lines:
                bbox = dummy_draw.textbbox((0, 0), line, font=font)
                width = bbox[2] - bbox[0]
                height = bbox[3] - bbox[1]
                line_sizes.append((width, height, bbox))
            max_width = max([w for w, h, _ in line_sizes]) + 2 * padding
            total_height = sum([h for w, h, _ in line_sizes]) + 2 * padding + (len(lines)-1)*2*scale
            img = Image.new("RGBA", (max_width, total_height), (255, 255, 255, 0))
            draw = ImageDraw.Draw(img)
            y = padding
            for i, line in enumerate(lines):
                w, h, bbox = line_sizes[i]
                offset_x = (max_width - w) // 2
                draw.text((offset_x, y - bbox[1]), line, font=font, fill=color)
                y += h + 2*scale
            return img

        DEFAULT_SIGNATURE_HEIGHT = 50
        DEFAULT_COMMENT_FONT_SIZE = 18
        font_path = os.path.join(os.path.dirname(__file__), "fonts", "THSarabunNew.ttf")
        
        if not os.path.isfile(font_path):
            return jsonify({'error': f"Font file not found: {font_path}"}), 500

        # เปิด PDF ที่เพิ่งสร้าง
        main_pdf = fitz.open(tmp_pdf)

        from collections import defaultdict
        sig_dict = defaultdict(list)
        for sig in signatures:
            page_number = int(sig.get('page', 0))
            x = int(sig['x'])
            y = int(sig['y'])
            # แยกประเภท PDF ที่จะแปะลายเซ็น
            pdf_type = sig.get('pdf_type', 'main')  # 'main' หรือ 'attachment'
            sig_dict[(pdf_type, page_number, x, y)].append(sig)

        # ฟังก์ชันสำหรับวาดลายเซ็น
        def process_signatures_on_pdf(pdf, sig_dict_filtered):
            for (page_number, x, y), sigs in sig_dict_filtered.items():
                if page_number >= len(pdf):
                    continue  # ข้ามถ้าหน้าไม่มี
                page = pdf[page_number]
                current_y = y
                has_lines = any('lines' in sig for sig in sigs)
                
                if has_lines:
                    for sig in sigs:
                        lines = sig.get('lines')
                        if not lines:
                            # fallback to old logic
                            if sig['type'] == 'text':
                                text = to_thai_digits(sig.get('text', ''))
                                font_size = 20 if sig.get('type') == 'comment' else DEFAULT_COMMENT_FONT_SIZE
                                font_weight = "bold" if sig.get('type') == 'comment' else "regular"
                                orig_color = sig.get('color', (2, 53, 139))
                                if isinstance(orig_color, (list, tuple)):
                                    r = min(int(orig_color[0]*0.8), 255)
                                    g = min(int(orig_color[1]*0.8), 255)
                                    b = min(int(orig_color[2]*0.8), 255)
                                    color = (r, g, b)
                                else:
                                    color = (2, 53, 139)
                                img = draw_text_image_v2(text, font_path, font_size=font_size, color=color, scale=1, font_weight=font_weight)
                                img_byte_arr = io.BytesIO()
                                img.save(img_byte_arr, format='PNG')
                                left_x = x - img.width // 2
                                rect = fitz.Rect(left_x, current_y, left_x + img.width, current_y + img.height)
                                page.insert_image(rect, stream=img_byte_arr.getvalue())
                                current_y += img.height
                            elif sig['type'] == 'image':
                                file_key = sig['file_key']
                                if file_key in request.files:
                                    signature_file = request.files[file_key]
                                    img = Image.open(signature_file)
                                    fixed_height = DEFAULT_SIGNATURE_HEIGHT
                                    ratio = fixed_height / img.height
                                    new_width = int(img.width * ratio)
                                    img = img.resize((new_width, fixed_height), resample=Image.LANCZOS)
                                    img_byte_arr = io.BytesIO()
                                    img.save(img_byte_arr, format='PNG')
                                    left_x = x - new_width // 2
                                    rect = fitz.Rect(left_x, current_y, left_x + new_width, current_y + fixed_height)
                                    page.insert_image(rect, stream=img_byte_arr.getvalue())
                                    current_y += fixed_height
                        else:
                            # draw lines in order
                            for line in lines:
                                line_type = line.get('type')
                                if line_type == 'image':
                                    file_key = line.get('file_key')
                                    if file_key and file_key in request.files:
                                        signature_file = request.files[file_key]
                                        img = Image.open(signature_file)
                                        fixed_height = DEFAULT_SIGNATURE_HEIGHT
                                        ratio = fixed_height / img.height
                                        new_width = int(img.width * ratio)
                                        img = img.resize((new_width, fixed_height), resample=Image.LANCZOS)
                                        img_byte_arr = io.BytesIO()
                                        img.save(img_byte_arr, format='PNG')
                                        left_x = x - new_width // 2
                                        rect = fitz.Rect(left_x, current_y, left_x + new_width, current_y + fixed_height)
                                        page.insert_image(rect, stream=img_byte_arr.getvalue())
                                        current_y += fixed_height
                                else:
                                    text_value = line.get('text') or line.get('value') or line.get('comment') or ''

                                    # ฟังก์ชันนับตัวอักษรที่มองเห็น (ไม่รวม tone marks, vowel marks)
                                    def count_visible_chars(s):
                                        thai_marks = set([
                                            '\u0E31', '\u0E34', '\u0E35', '\u0E36', '\u0E37',
                                            '\u0E38', '\u0E39', '\u0E3A', '\u0E47', '\u0E48',
                                            '\u0E49', '\u0E4A', '\u0E4B', '\u0E4C', '\u0E4D', '\u0E4E'
                                        ])
                                        return len([c for c in s if c not in thai_marks])

                                    # ฟังก์ชันตัดข้อความตามจำนวนตัวอักษรที่มองเห็น
                                    def wrap_by_visible_chars(text, max_chars=30):
                                        if count_visible_chars(text) <= max_chars:
                                            return [text]

                                        lines = []
                                        current = ""
                                        for char in text:
                                            test = current + char
                                            if count_visible_chars(test) <= max_chars:
                                                current = test
                                            else:
                                                if current:
                                                    lines.append(current)
                                                current = char
                                        if current:
                                            lines.append(current)
                                        return lines

                                    # ถ้าเป็น comment และมี - ให้แยกเป็นหลายบรรทัด
                                    if line_type == 'comment' and '-' in text_value:
                                        # แยกตาม - และเก็บ - ไว้ข้างหน้าแต่ละบรรทัด
                                        parts = text_value.split('-')
                                        # เอา parts ที่ไม่ว่างเปล่าออกมา
                                        text_lines = ['-' + part for part in parts if part.strip()]

                                        # ตัดแต่ละบรรทัดให้ไม่เกิน 15 ตัวอักษรที่มองเห็น
                                        wrapped_lines = []
                                        for tline in text_lines:
                                            wrapped = wrap_by_visible_chars(tline, max_chars=15)
                                            wrapped_lines.extend(wrapped)

                                        text = '\n'.join([to_thai_digits(t) for t in wrapped_lines])
                                    else:
                                        text = to_thai_digits(text_value)

                                    font_size = 20 if line_type == 'comment' else DEFAULT_COMMENT_FONT_SIZE
                                    font_weight = "bold" if line_type == 'comment' else "regular"
                                    orig_color = line.get('color', (2, 53, 139))
                                    if isinstance(orig_color, (list, tuple)):
                                        r = min(int(orig_color[0]*0.8), 255)
                                        g = min(int(orig_color[1]*0.8), 255)
                                        b = min(int(orig_color[2]*0.8), 255)
                                        color = (r, g, b)
                                    else:
                                        color = (2, 53, 139)
                                    img = draw_text_image_v2(text, font_path, font_size=font_size, color=color, scale=1, font_weight=font_weight)
                                    img_byte_arr = io.BytesIO()
                                    img.save(img_byte_arr, format='PNG')
                                    left_x = x - img.width // 2
                                    rect = fitz.Rect(left_x, current_y, left_x + img.width, current_y + img.height)
                                    page.insert_image(rect, stream=img_byte_arr.getvalue())
                                    current_y += img.height
                else:
                    # fallback to old logic
                    sigs_sorted = sorted(sigs, key=lambda s: 0 if s['type'] == 'text' else 1)
                    for sig in sigs_sorted:
                        if sig['type'] == 'text':
                            text = to_thai_digits(sig.get('text', ''))
                            font_size = 20 if sig.get('type') == 'comment' else DEFAULT_COMMENT_FONT_SIZE
                            font_weight = "bold" if sig.get('type') == 'comment' else "regular"
                            orig_color = sig.get('color', (2, 53, 139))
                            if isinstance(orig_color, (list, tuple)):
                                r = min(int(orig_color[0]*0.8), 255)
                                g = min(int(orig_color[1]*0.8), 255)
                                b = min(int(orig_color[2]*0.8), 255)
                                color = (r, g, b)
                            else:
                                color = (2, 53, 139)
                            img = draw_text_image_v2(text, font_path, font_size=font_size, color=color, scale=1, font_weight=font_weight)
                            img_byte_arr = io.BytesIO()
                            img.save(img_byte_arr, format='PNG')
                            left_x = x - img.width // 2
                            rect = fitz.Rect(left_x, current_y, left_x + img.width, current_y + img.height)
                            page.insert_image(rect, stream=img_byte_arr.getvalue())
                            current_y += img.height
                        elif sig['type'] == 'image':
                            file_key = sig['file_key']
                            if file_key in request.files:
                                signature_file = request.files[file_key]
                                img = Image.open(signature_file)
                                fixed_height = DEFAULT_SIGNATURE_HEIGHT
                                ratio = fixed_height / img.height
                                new_width = int(img.width * ratio)
                                img = img.resize((new_width, fixed_height), resample=Image.LANCZOS)
                                img_byte_arr = io.BytesIO()
                                img.save(img_byte_arr, format='PNG')
                                left_x = x - new_width // 2
                                rect = fitz.Rect(left_x, current_y, left_x + new_width, current_y + fixed_height)
                                page.insert_image(rect, stream=img_byte_arr.getvalue())
                                current_y += fixed_height

        # แยกลายเซ็นตามประเภท PDF
        main_sigs = {(page, x, y): sigs for (pdf_type, page, x, y), sigs in sig_dict.items() if pdf_type == 'main'}
        attachment_sigs = {(page, x, y): sigs for (pdf_type, page, x, y), sigs in sig_dict.items() if pdf_type == 'attachment'}

        # วาดลายเซ็นลง PDF หลัก
        if main_sigs:
            process_signatures_on_pdf(main_pdf, main_sigs)

        # วาดลายเซ็นลงเอกสารแนบ (ถ้ามี)
        if attachment_pdf and attachment_sigs:
            process_signatures_on_pdf(attachment_pdf, attachment_sigs)

        # รวม PDF ทั้งหมดเป็นไฟล์เดียว
        final_pdf = fitz.open()
        
        # เพิ่มหน้าจาก main PDF
        final_pdf.insert_pdf(main_pdf)
        
        # เพิ่มหน้าจาก attachment PDF (ถ้ามี)
        if attachment_pdf:
            final_pdf.insert_pdf(attachment_pdf)

        # บันทึก PDF ที่มีลายเซ็นแล้ว
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as final_pdf_file:
            final_pdf.save(final_pdf_file.name)
        
        # ปิด PDF ทั้งหมด
        main_pdf.close()
        if attachment_pdf:
            attachment_pdf.close()
        final_pdf.close()
        
        # ลบไฟล์ชั่วคราว
        os.unlink(tmp_pdf)
        
        return send_file(final_pdf_file.name, mimetype="application/pdf", as_attachment=True, download_name="signed_memo.pdf")
        
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({'error': str(e)}), 500


@app.route('/PDFmerge', methods=['POST'])
def merge_pdfs():
    """รวมไฟล์ PDF 2 ไฟล์เป็นไฟล์เดียว"""
    try:
        # ตรวจสอบไฟล์ที่ส่งมา
        if 'pdf1' not in request.files:
            return jsonify({'error': 'No pdf1 file uploaded'}), 400
        if 'pdf2' not in request.files:
            return jsonify({'error': 'No pdf2 file uploaded'}), 400
        
        pdf1_file = request.files['pdf1']
        pdf2_file = request.files['pdf2']
        
        # อ่านไฟล์ PDF เป็น bytes (blob)
        pdf1_bytes = pdf1_file.read()
        pdf2_bytes = pdf2_file.read()
        
        # เปิดไฟล์ PDF จาก bytes
        pdf1 = fitz.open(stream=pdf1_bytes, filetype="pdf")
        pdf2 = fitz.open(stream=pdf2_bytes, filetype="pdf")
        
        # สร้าง PDF ใหม่สำหรับรวมไฟล์
        merged_pdf = fitz.open()
        
        # เพิ่มหน้าจาก PDF แรก
        merged_pdf.insert_pdf(pdf1)
        
        # เพิ่มหน้าจาก PDF ที่สอง
        merged_pdf.insert_pdf(pdf2)
        
        # ปิดไฟล์ต้นฉบับ
        pdf1.close()
        pdf2.close()
        
        # บันทึกไฟล์ที่รวมแล้ว
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as merged_file:
            merged_pdf.save(merged_file.name)
        
        merged_pdf.close()
        
        # ส่งไฟล์กลับ
        return send_file(merged_file.name, mimetype="application/pdf", as_attachment=True, download_name="merged.pdf")
        
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/receive_num', methods=['POST'])
def receive_num():
    """
    multipart/form-data:
      - pdf: ไฟล์ PDF  
      - payload: JSON string:
        {
          "page": 0,
          "color": [2,53,139],
          "register_no": "2567/506",
          "date": "20 ก.ย. 67",
          "time": "10.30 น.",
          "receiver": "ดวงดี"
        }
    หมายเหตุ: ตรายางจะวาดที่มุมขวาบนแบบ fix ไม่ต้องส่ง x,y
    """
    print("[DEBUG] /receive_num API called")
    try:
        if 'pdf' not in request.files:
            return jsonify({'error': 'No PDF file uploaded'}), 400
        if 'payload' not in request.form:
            return jsonify({'error': 'No payload'}), 400

        p = json.loads(request.form['payload'])
        print(f"[DEBUG] Payload received: {p}")
        page_no = int(p.get('page', 0))
        color = tuple(p.get('color', [2,53,139]))
        print(f"[DEBUG] Page: {page_no}, color: {color}")

        # เปิด PDF
        pdf_bytes = request.files['pdf'].read()
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        if page_no >= len(doc):
            return jsonify({'error': 'Page out of range'}), 400
        page = doc[page_no]

        # ฟอนต์
        font_path = os.path.join(os.path.dirname(__file__), "fonts", "THSarabunNew.ttf")
        bold_font_path = os.path.join(os.path.dirname(__file__), "fonts", "THSarabunNew Bold.ttf")
        print(f"[DEBUG] Font paths: {font_path}, {bold_font_path}")
        print(f"[DEBUG] Font exists: regular={os.path.isfile(font_path)}, bold={os.path.isfile(bold_font_path)}")
        if not os.path.isfile(font_path) or not os.path.isfile(bold_font_path):
            return jsonify({'error': 'THSarabunNew fonts not found'}), 500


        # Fix ตำแหน่งที่มุมขวาบนของกระดาษ
        page_w = page.rect.width
        page_h = page.rect.height
        
        # คำนวณตำแหน่งมุมขวาบน (เว้นขอบเล็กน้อย)
        margin = 20  # เว้นขอบ 20 pixel
        frame_width = 200
        frame_height = 80
        
        # ตำแหน่งกึ่งกลางตรายาง = มุมขวาบน - ขอบ - ครึ่งตรายาง
        center_x = page_w - margin - frame_width//2
        center_y = margin + frame_height//2
        
        print(f"[DEBUG] Page size: {page_w}x{page_h}")
        print(f"[DEBUG] Fixed stamp position: center_x={center_x}, center_y={center_y}")
        
        # วาดกรอบสี่เหลี่ยมสีน้ำเงิน (เหมือนตรายาง)
        box_left = center_x - frame_width//2
        box_top = center_y - frame_height//2
        box_right = center_x + frame_width//2
        box_bottom = center_y + frame_height//2
        
        box_rect = fitz.Rect(box_left, box_top, box_right, box_bottom)
        # ใช้สีเดียวกับฟอนต์ (สีน้ำเงินอ่อนลง)
        box_color = (color[0]/255, color[1]/255, color[2]/255)  # แปลง RGB เป็น 0-1
        page.draw_rect(box_rect, color=box_color, width=2)  # กรอบสีเดียวกับฟอนต์ หนา 1px
        print(f"[DEBUG] Drew blue frame at {box_rect} with color {box_color}")

        # *** ลบ test text ออก และใช้ตัวอย่างง่ายๆ ***
        
        # เส้นหัวข้อกรอบตรา 4 บรรทัด (หนา) - commented out for testing
        """
        header_lines = [
            "ศูนย์การศึกษาพิเศษ เขตการศึกษา ๖ จ.ลพบุรี",
            "เลขทะเบียนรับที่ ..........",
            "วันที่ ........../............/............ เวลา ..........น.",
            "ผู้รับ ........................"
        ]
        gap = int(bh/4)  # ระยะห่างแต่ละบรรทัดในกรอบ

        start_y = center_y - ( (len(header_lines)-1) * gap // 2 )
        print(f"[DEBUG] start_y: {start_y}, gap: {gap}")
        for i, text in enumerate(header_lines):
            print(f"[DEBUG] Drawing header line {i}: {text}")
            img = draw_text_img(text, size=16, bold=True)
            y_pos = start_y + i*gap
            print(f"[DEBUG] Position: center_x={center_x}, y={y_pos}")
            paste_center(img, center_x, y_pos)
        """

        # วาดข้อมูลตรา - ใช้ PyMuPDF text แทน PIL
        register_no = p.get('register_no','')
        date_text = p.get('date','')
        time_text = p.get('time','')
        receiver_text = p.get('receiver','')
        
        print(f"[DEBUG] Data to insert: register_no='{register_no}', date='{date_text}', time='{time_text}', receiver='{receiver_text}'")
        
        # ใช้วิธีเดียวกับ endpoint อื่น - draw_text_image + insert_image
        def draw_text_img(text, size=18, bold=False):
            fp = bold_font_path if bold else font_path
            print(f"[DEBUG] Creating text image: '{text}', size={size}, bold={bold}, font={fp}")
            img = draw_text_image(to_thai_digits(text), fp, size, color, scale=1)
            print(f"[DEBUG] Text image created: {img.width}x{img.height}")
            return img

        def paste_center(img, center_x, center_y):
            left = center_x - img.width//2
            top  = center_y - img.height//2
            rect = fitz.Rect(left, top, left+img.width, top+img.height)
            print(f"[DEBUG] Pasting image at rect: {rect}, image size: {img.width}x{img.height}")
            bio = io.BytesIO(); img.save(bio, format='PNG')
            page.insert_image(rect, stream=bio.getvalue())
            print(f"[DEBUG] Image inserted successfully")
        
        # เส้นหัวข้อกรอบตรา 4 บรรทัด (หนา)
        header_lines = [
            "ศูนย์การศึกษาพิเศษ เขตการศึกษา ๖ จ.ลพบุรี",
            f"เลขทะเบียนรับที่ {register_no}",
            f"วันที่ {date_text} เวลา {time_text}",
            f"ผู้รับ {receiver_text}"
        ]
        gap = 16  # ระยะห่างแต่ละบรรทัดใหม่ (แทนที่ bh/4 ที่มากเกินไป)

        start_y = center_y - ( (len(header_lines)-1) * gap // 2 )
        print(f"[DEBUG] start_y: {start_y}, gap: {gap}")
        for i, text in enumerate(header_lines):
            print(f"[DEBUG] Drawing header line {i}: {text}")
            img = draw_text_img(text, size=16, bold=True)
            y_pos = start_y + i*gap
            print(f"[DEBUG] Position: center_x={center_x}, y={y_pos}")
            paste_center(img, center_x, y_pos)

        # ส่งไฟล์กลับ
        print("[DEBUG] Saving final PDF...")
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as outpdf:
            doc.save(outpdf.name)
        doc.close()
        print(f"[DEBUG] PDF saved, sending response...")
        
        # เพิ่ม debug response header
        response = send_file(outpdf.name, mimetype="application/pdf", as_attachment=True, download_name="receive_num.pdf")
        response.headers['X-Debug'] = 'receive_num_processed'
        return response

    except Exception as e:
        print(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/stamp_summary', methods=['POST'])
def stamp_summary():
    """
    multipart/form-data:
      - pdf: ไฟล์ PDF ต้นฉบับ
      - sign_png: ไฟล์ลายเซ็นธุรการ (PNG โปร่งใส)
      - payload: JSON string:
        {
          "summary": "เรื่อง การขออนุมัติโครงการ",
          "group_name": "กลุ่มวิชาการ",
          "receiver_name": "นายสมชาย รับผิดชอบ",
          "date": "25 ก.ย. 67",
          "page": 0,  // optional, default = 0
          "x": 100,   // optional, center x position
          "y": 200,   // optional, center y position
          "width": 200,  // optional, stamp width (for positioning)
          "height": 150  // optional, stamp height (for positioning)
        }

    ถ้าไม่ระบุ x, y จะวาดที่มุมซ้ายล่าง (default)
    ถ้าระบุ x, y จะใช้เป็น center position แบบเดียวกับลายเซ็น
    """
    print("[DEBUG] /stamp_summary API called")
    try:
        # ตรวจสอบไฟล์และข้อมูลที่ส่งมา
        if 'pdf' not in request.files:
            return jsonify({'error': 'No PDF file uploaded'}), 400
        if 'sign_png' not in request.files:
            return jsonify({'error': 'No signature PNG file uploaded'}), 400
        if 'payload' not in request.form:
            return jsonify({'error': 'No payload'}), 400

        # อ่านข้อมูล
        pdf_file = request.files['pdf']
        sign_file = request.files['sign_png']

        p = json.loads(request.form['payload'])
        summary = p.get('summary', '')
        group_name = p.get('group_name', '')
        receiver_name = p.get('receiver_name', '')
        date = p.get('date', '')

        # รองรับการระบุตำแหน่ง
        page_number = int(p.get('page', 0))
        pos_x = p.get('x', None)  # center x
        pos_y = p.get('y', None)  # center y
        pos_width = p.get('width', None)  # width for positioning
        pos_height = p.get('height', None)  # height for positioning

        print(f"[DEBUG] Data: summary={summary}, group={group_name}, receiver={receiver_name}, date={date}")
        print(f"[DEBUG] Position: page={page_number}, x={pos_x}, y={pos_y}, width={pos_width}, height={pos_height}")

        # เปิด PDF
        pdf_bytes = pdf_file.read()
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        page = doc[page_number]  # ใช้หน้าที่ระบุ

        # ฟอนต์ไทย
        font_path = os.path.join(os.path.dirname(__file__), "fonts", "THSarabunNew.ttf")
        bold_font_path = os.path.join(os.path.dirname(__file__), "fonts", "THSarabunNew Bold.ttf")
        if not os.path.isfile(font_path) or not os.path.isfile(bold_font_path):
            return jsonify({'error': 'THSarabunNew fonts not found'}), 500

        # เตรียมข้อมูลสำหรับคำนวณความสูง
        page_w = page.rect.width
        page_h = page.rect.height
        margin = 30
        stamp_width = 200
        
        # ข้อมูลที่จะแสดงในตรา
        header_text = "เรียน ผอ. ศกศ.เขต ๖ จ.ลพบุรี"
        signature_text = "ลงชื่อ"
        receiver_text = f"ผู้รับ  {receiver_name}"
        date_text = f"วันที่ {date}"
        
        # สร้างฟังก์ชันวาดข้อความ (รองรับการ wrap text)
        def draw_text_img(text, size=16, bold=False, max_width=None):
            from PIL import ImageFont, ImageDraw, Image
            fp = bold_font_path if bold else font_path
            color_rgb = (2, 53, 139)  # สีน้ำเงิน
            text = to_thai_digits(text)

            # ถ้าไม่มี max_width ใช้วิธีเดิม
            if max_width is None:
                img = draw_text_image(text, fp, size, color_rgb, scale=1)
                return img

            # มี max_width ให้ wrap text
            font = ImageFont.truetype(fp, size)
            padding = 4

            # แยกข้อความเป็นบรรทัดตาม max_width
            words = text.split(' ')
            lines = []
            current_line = ""

            for i, word in enumerate(words):
                test_line = current_line + (" " if current_line else "") + word
                bbox = font.getbbox(test_line)
                text_width = bbox[2] - bbox[0]

                if text_width <= max_width - 2 * padding:
                    current_line = test_line
                else:
                    # กรณีพิเศษ: ถ้าบรรทัดปัจจุบันคือ "เรื่อง" หรือ "เห็นควรมอบ" ต้องเก็บคำถัดไปไว้ด้วยเสมอ
                    if current_line == "เรื่อง" or current_line == "เห็นควรมอบ":
                        # บังคับรวม prefix กับคำแรกของข้อความ แล้วตัดทีละตัวอักษร
                        # เพิ่ม space ครั้งเดียวหลัง prefix
                        current_line = current_line + " "
                        for char in word:
                            test_char = current_line + char
                            char_bbox = font.getbbox(test_char)
                            char_width = char_bbox[2] - char_bbox[0]

                            if char_width <= max_width - 2 * padding:
                                current_line = test_char
                            else:
                                if current_line:
                                    lines.append(current_line)
                                current_line = char
                    else:
                        # ถ้าคำยาวเกินไป ต้องตัดทีละตัวอักษร
                        word_bbox = font.getbbox(word)
                        word_width = word_bbox[2] - word_bbox[0]

                        if word_width > max_width - 2 * padding:
                            # คำยาวเกิน ต้องตัดทีละตัวอักษร
                            if current_line:
                                lines.append(current_line)
                                current_line = ""

                            # ตัดคำยาวทีละตัวอักษร
                            for char in word:
                                test_char = current_line + char
                                char_bbox = font.getbbox(test_char)
                                char_width = char_bbox[2] - char_bbox[0]

                                if char_width <= max_width - 2 * padding:
                                    current_line += char
                                else:
                                    if current_line:
                                        lines.append(current_line)
                                    current_line = char
                        else:
                            # คำไม่ยาวเกิน แต่รวมกับบรรทัดปัจจุบันแล้วยาวเกิน
                            if current_line:
                                lines.append(current_line)
                            current_line = word

            if current_line:
                lines.append(current_line)

            # คำนวณขนาดภาพ
            max_line_width = 0
            fixed_line_height = 14  # ใช้ความสูงคงที่

            for line in lines:
                bbox = font.getbbox(line)
                line_width = bbox[2] - bbox[0]
                max_line_width = max(max_line_width, line_width)

            img_width = max_line_width + 2 * padding
            # ใช้ความสูงคงที่ 14px ต่อบรรทัด
            img_height = len(lines) * fixed_line_height + 2 * padding

            # สร้างภาพ
            img = Image.new("RGBA", (img_width, img_height), (255, 255, 255, 0))
            draw = ImageDraw.Draw(img)

            y = padding
            for i, line in enumerate(lines):
                draw.text((padding, y), line, font=font, fill=color_rgb)
                # ใช้ความสูงคงที่
                y += fixed_line_height

            # เก็บจำนวนบรรทัดไว้ใน attribute ของภาพ
            img.line_count = len(lines)
            return img

        
        # กำหนด debug_logs ในระดับท้องถิ่น
        debug_logs = []
        
        # ฟังก์ชันตัดข้อความแบบง่าย - 30 ตัวอักษรต่อบรรทัด (ไม่นับ vowel/tone marks)
        def wrap_text(text, max_chars_approx):
            # ใช้ debug_logs ในระดับท้องถิ่นแทน global
            try:
                debug_logs.append(f"wrap_text called: text='{text[:20]}...', max_chars_approx={max_chars_approx}")
            except:
                pass  # ไม่ให้ error หยุดการทำงาน
            # ตัวอักษรไทยที่ไม่ควรนับ (vowel marks, tone marks)
            thai_marks = set([
                '\u0E31',  # ั
                '\u0E34',  # ิ
                '\u0E35',  # ี
                '\u0E36',  # ึ
                '\u0E37',  # ื
                '\u0E38',  # ุ
                '\u0E39',  # ู
                '\u0E3A',  # ฺ
                '\u0E47',  # ็
                '\u0E48',  # ่
                '\u0E49',  # ้
                '\u0E4A',  # ๊
                '\u0E4B',  # ๋
                '\u0E4C',  # ์
                '\u0E4D',  # ํ
                '\u0E4E'   # ๎
            ])
            
            def count_visible_chars(s):
                """นับตัวอักษรที่มองเห็นได้ (ไม่รวม marks)"""
                count = len([c for c in s if c not in thai_marks])
                try:
                    debug_logs.append(f"count_visible_chars('{s[:20]}...') = {count}")
                except:
                    pass
                return count
            
            def cut_at_visible_chars(s, max_visible):
                """ตัดข้อความที่ตัวอักษรที่มองเห็นได้ โดยคำต้องไม่ขาด"""
                if count_visible_chars(s) <= max_visible:
                    return s
                
                visible_count = 0
                
                for i, char in enumerate(s):
                    if char not in thai_marks:
                        visible_count += 1
                        
                        # ถ้าถึง 30 ตัวแล้ว
                        if visible_count >= max_visible:
                            current_pos = i + 1
                            
                            # หาตำแหน่งที่ไม่ทำให้คำขาด (หาช่องว่างหรือจุดสิ้นสุดถัดไป)
                            for j in range(current_pos, len(s)):
                                if s[j] == ' ' or s[j] in '.!?,:;':
                                    return s[:j]
                            
                            # ถ้าไม่เจอจุดแบ่งคำ ให้หาตำแหน่งย้อนกลับที่ไม่ใช่ Thai mark
                            cut_pos = current_pos
                            while cut_pos > 0 and s[cut_pos-1] in thai_marks:
                                cut_pos -= 1
                            
                            return s[:max(cut_pos, current_pos-5)]  # อย่างน้อยตัดไม่เกิน 5 ตัวย้อนกลับ
                
                return s
            
            lines = []
            max_chars = 30
            
            # ถ้าขึ้นต้นด้วย "เรื่อง " ให้รักษาไว้กับคำถัดไป
            if text.startswith("เรื่อง "):
                words = text.split(' ')
                if len(words) >= 2:
                    first_group = f"เรื่อง {words[1]}"
                    rest_text = ' '.join(words[2:])
                    
                    # ถ้า first_group เกิน 30 ตัวอักษรที่มองเห็น ก็ตัดแบบธรรมดา
                    if count_visible_chars(first_group) > max_chars:
                        text = text  # ใช้วิธีตัดธรรมดา
                    else:
                        # ใส่ first_group ในบรรทัดแรก
                        lines.append(first_group)
                        text = rest_text
            
            # ตัดที่ 30 ตัวอักษรที่มองเห็นได้
            visible_count = count_visible_chars(text)
            try:
                debug_logs.append(f"Before while: visible_chars={visible_count}, max_chars={max_chars}")
                debug_logs.append(f"Should enter while loop? {visible_count > max_chars}")
            except:
                pass
            
            while count_visible_chars(text) > max_chars:
                try:
                    debug_logs.append("Entering while loop, cutting text...")
                except:
                    pass
                cut_text = cut_at_visible_chars(text, max_chars)
                lines.append(cut_text)
                text = text[len(cut_text):]
                try:
                    debug_logs.append(f"Added line: '{cut_text[:20]}...', remaining: '{text[:20]}...'")
                    debug_logs.append(f"Remaining text visible chars: {count_visible_chars(text)}")
                except:
                    pass
            
            if text.strip():
                lines.append(text)
                try:
                    debug_logs.append(f"Added final line: '{text[:20]}...'")
                except:
                    pass
            
            try:
                debug_logs.append(f"Final result: {len(lines)} lines")
            except:
                pass
            return lines

        # สร้างข้อความทั้งหมดก่อนเพื่อคำนวณความสูงจริง
        font_size = 16
        text_max_width = stamp_width - 20  # เว้นขอบซ้าย-ขวา 10 px

        # สร้างภาพข้อความทั้งหมดก่อน
        text1 = "เรียน ผอ. ศกศ.เขต ๖ จ.ลพบุรี"
        img1 = draw_text_img(text1, size=font_size, bold=True, max_width=text_max_width)

        # ฟังก์ชันสร้างภาพข้อความแบบ mixed (prefix ตัวหนา, content ตัวปกติ)
        def draw_mixed_text_img(prefix, content, size=16, max_width=None):
            from PIL import ImageFont, ImageDraw, Image
            color_rgb = (2, 53, 139)
            padding = 4
            fixed_line_height = 14

            # แปลงเป็นเลขไทย
            prefix = to_thai_digits(prefix)
            content = to_thai_digits(content)

            bold_font = ImageFont.truetype(bold_font_path, size)
            normal_font = ImageFont.truetype(font_path, size)

            # วาด prefix (ตัวหนา) + content (ตัวปกติ) โดย wrap ทั้งหมด
            full_text = prefix + " " + content

            # แยกข้อความเป็นบรรทัดตาม max_width (ใช้โค้ดเดียวกับ draw_text_img)
            words = full_text.split(' ')
            lines = []
            current_line = ""

            for i, word in enumerate(words):
                # ใช้ bold font สำหรับคำแรก (prefix), normal font สำหรับที่เหลือ
                test_font = bold_font if i == 0 else normal_font
                test_line = current_line + (" " if current_line else "") + word
                bbox = test_font.getbbox(test_line) if i == 0 else normal_font.getbbox(test_line)
                text_width = bbox[2] - bbox[0]

                if text_width <= max_width - 2 * padding:
                    current_line = test_line
                else:
                    # กรณีพิเศษ: ถ้าบรรทัดปัจจุบันเป็น prefix
                    if current_line == prefix:
                        current_line = current_line + " "
                        for char in word:
                            test_char = current_line + char
                            char_bbox = normal_font.getbbox(test_char)
                            char_width = char_bbox[2] - char_bbox[0]

                            if char_width <= max_width - 2 * padding:
                                current_line = test_char
                            else:
                                if current_line:
                                    lines.append(current_line)
                                current_line = char
                    else:
                        word_bbox = normal_font.getbbox(word)
                        word_width = word_bbox[2] - word_bbox[0]

                        if word_width > max_width - 2 * padding:
                            if current_line:
                                lines.append(current_line)
                                current_line = ""

                            for char in word:
                                test_char = current_line + char
                                char_bbox = normal_font.getbbox(test_char)
                                char_width = char_bbox[2] - char_bbox[0]

                                if char_width <= max_width - 2 * padding:
                                    current_line += char
                                else:
                                    if current_line:
                                        lines.append(current_line)
                                    current_line = char
                        else:
                            if current_line:
                                lines.append(current_line)
                            current_line = word

            if current_line:
                lines.append(current_line)

            # คำนวณความกว้างสูงสุด
            max_line_width = 0
            for line in lines:
                # ใช้ bold font สำหรับบรรทัดแรกที่มี prefix
                if line.startswith(prefix):
                    # วัดความกว้างแบบ mixed
                    prefix_bbox = bold_font.getbbox(prefix)
                    prefix_width = prefix_bbox[2] - prefix_bbox[0]
                    rest = line[len(prefix):]
                    rest_bbox = normal_font.getbbox(rest)
                    rest_width = rest_bbox[2] - rest_bbox[0]
                    line_width = prefix_width + rest_width
                else:
                    bbox = normal_font.getbbox(line)
                    line_width = bbox[2] - bbox[0]
                max_line_width = max(max_line_width, line_width)

            img_width = max_line_width + 2 * padding
            img_height = len(lines) * fixed_line_height + 2 * padding

            # สร้างภาพและวาดข้อความ
            img = Image.new("RGBA", (img_width, img_height), (255, 255, 255, 0))
            draw = ImageDraw.Draw(img)

            y = padding
            for line in lines:
                if line.startswith(prefix):
                    # วาด prefix ตัวหนา
                    draw.text((padding, y), prefix, font=bold_font, fill=color_rgb)
                    prefix_bbox = bold_font.getbbox(prefix)
                    prefix_width = prefix_bbox[2] - prefix_bbox[0]
                    # วาดส่วนที่เหลือตัวปกติ
                    rest = line[len(prefix):]
                    draw.text((padding + prefix_width, y), rest, font=normal_font, fill=color_rgb)
                else:
                    # วาดตัวปกติทั้งหมด
                    draw.text((padding, y), line, font=normal_font, fill=color_rgb)
                y += fixed_line_height

            img.line_count = len(lines)
            return img

        img_subject = draw_mixed_text_img("เรื่อง", summary, size=font_size, max_width=text_max_width)
        img_assign = draw_mixed_text_img("เห็นควรมอบ", group_name, size=font_size, max_width=text_max_width)

        sign_img_temp = Image.open(sign_file)
        sign_height = 30
        ratio = sign_height / sign_img_temp.height
        sign_width = int(sign_img_temp.width * ratio)
        sign_img = sign_img_temp.resize((sign_width, sign_height), Image.LANCZOS)

        sign_text = "ลงชื่อ"
        img_sign_text = draw_text_img(sign_text, size=font_size, bold=False)

        receiver_text_str = f"ผู้รับ  {receiver_name}"
        img_receiver = draw_text_img(receiver_text_str, size=font_size, bold=False)

        date_text_str = f"วันที่ {date}"
        img_date = draw_text_img(date_text_str, size=font_size, bold=False)

        # คำนวณความสูงจริงตามที่จะใช้ในการวาด
        padding_top = 8
        padding_bottom = 8
        line_height = 14

        total_height = padding_top
        total_height += line_height  # เรียน ผอ.
        total_height += img_subject.height + 2  # เรื่อง (ใช้ความสูงจริง)
        total_height += img_assign.height + 2  # เห็นควรมอบ (ใช้ความสูงจริง)
        total_height += line_height + 2  # ลงชื่อ
        total_height += line_height  # ผู้รับ
        total_height += line_height  # วันที่
        total_height += padding_bottom

        stamp_height = int(total_height)

        # คำนวณตำแหน่งกรอบ
        # ถ้าระบุ x, y ให้ใช้เป็น top-center position (x=center, y=top)
        if pos_x is not None and pos_y is not None:
            # Frontend ส่ง (x, y) มาโดยที่ x=center, y=top ของตรา
            # ต้องแปลง Y-axis จาก "บนเป็นล่าง" ให้เป็น "ล่างเป็นบน"
            center_x = int(pos_x)

            # แปลง Y จาก top position เป็น center position ใน PDF coordinate
            # adjusted_y = page_height - y + 30 (แปลง Y-axis)
            # center_y = adjusted_y + stamp_height/2 (เลื่อนลงครึ่งหนึ่งของความสูง)
            adjusted_y = page_h - pos_y + 30
            center_y = adjusted_y + stamp_height // 2

            print(f"[DEBUG] Using custom position (top-center): original=(x={pos_x}, y_top={pos_y})")
            print(f"[DEBUG] Y-axis flip: y_top={pos_y} -> adjusted_y_top={adjusted_y} -> center_y={center_y}")
            print(f"[DEBUG] Final center position: ({center_x}, {center_y}), stamp_height={stamp_height}")
        else:
            # ใช้ default position (มุมซ้ายล่าง)
            margin = 30
            center_x = margin + stamp_width//2
            center_y = page_h - margin - stamp_height//2
            print(f"[DEBUG] Using default position (bottom-left): center=({center_x}, {center_y})")

        # วาดกรอบตรา
        box_left = center_x - stamp_width//2
        box_top = center_y - stamp_height//2
        box_right = center_x + stamp_width//2
        box_bottom = center_y + stamp_height//2

        print(f"[DEBUG] Stamp box: left={box_left}, top={box_top}, right={box_right}, bottom={box_bottom}")
        print(f"[DEBUG] Stamp dimensions: {stamp_width}x{stamp_height}")

        box_rect = fitz.Rect(box_left, box_top, box_right, box_bottom)
        box_color = (2/255, 53/255, 139/255)
        page.draw_rect(box_rect, color=box_color, width=2)

        def paste_at_position(img, x, y):
            rect = fitz.Rect(x, y, x+img.width, y+img.height)
            bio = io.BytesIO()
            img.save(bio, format='PNG')
            page.insert_image(rect, stream=bio.getvalue())

        # วาดข้อความในตรา (ใช้ภาพที่สร้างไว้แล้ว)
        # ใช้ระยะห่างแบบเดิม
        line_height = 14
        current_y = box_top + padding_top

        # เรียน ผอ.
        paste_at_position(img1, box_left + 10, current_y)
        current_y += line_height

        # เรื่อง + summary
        paste_at_position(img_subject, box_left + 10, current_y)
        # ใช้ความสูงจริงของภาพ + ระยะห่างเล็กน้อย
        current_y += img_subject.height + 2

        # เห็นควรมอบ + group_name
        paste_at_position(img_assign, box_left + 10, current_y)
        current_y += img_assign.height + 2

        # ลายเซ็น (ใช้ภาพที่สร้างไว้แล้ว)
        center_x_frame = box_left + stamp_width//2

        # คำนวณตำแหน่งเริ่มต้นให้อยู่กึ่งกลาง
        total_width = img_sign_text.width + 5 + sign_width
        start_x = center_x_frame - total_width//2

        sign_y = current_y
        # วาง "ลงชื่อ" ก่อน
        paste_at_position(img_sign_text, start_x, sign_y)
        # วางลายเซ็นติดข้าง
        paste_at_position(sign_img, start_x + img_sign_text.width + 5, sign_y)

        current_y += line_height + 2

        # ผู้รับ (กึ่งกลาง - ใช้ภาพที่สร้างไว้แล้ว)
        paste_at_position(img_receiver, center_x_frame - img_receiver.width//2, current_y)
        current_y += line_height

        # วันที่ (กึ่งกลาง - ใช้ภาพที่สร้างไว้แล้ว)
        paste_at_position(img_date, center_x_frame - img_date.width//2, current_y)

        # ส่งไฟล์กลับ
        print("[DEBUG] Saving final PDF...")
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as outpdf:
            doc.save(outpdf.name)
            doc.close()
            print(f"[DEBUG] PDF saved, sending response...")
            
            response = send_file(outpdf.name, mimetype="application/pdf", as_attachment=True, download_name="summary_stamped.pdf")
            response.headers['X-Debug'] = 'stamp_summary_processed'
            return response

    except Exception as e:
        print(f"[ERROR] {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/add_signature_receive', methods=['POST'])
def add_signature_receive():
    """
    รวมการทำงานของ /add_signature_v2 และ /stamp_summary ในครั้งเดียว

    multipart/form-data:
      - pdf: ไฟล์ PDF ต้นฉบับ
      - signatures: JSON string สำหรับลายเซ็น (เหมือน /add_signature_v2)
      - signature_files: ไฟล์รูปลายเซ็นต่างๆ
      - sign_png: ไฟล์ลายเซ็นธุรการสำหรับตราสรุป (PNG โปร่งใส)
      - summary_payload: JSON string สำหรับตราสรุป:
        {
          "summary": "เรื่อง การขออนุมัติโครงการ",
          "group_name": "กลุ่มวิชาการ",
          "receiver_name": "นายสมชาย รับผิดชอบ",
          "date": "25 ก.ย. 67"
        }
    """
    print("[DEBUG] /add_signature_receive API called")

    # ฟังก์ชันวาดข้อความเป็นภาพ (v2)
    def draw_text_image_v2(text, font_path, font_size=20, color=(2, 53, 139), scale=1, font_weight="regular"):
        from PIL import ImageFont, ImageDraw, Image
        # เลือก font ตาม font_weight
        if font_weight == "bold":
            font_path = os.path.join(os.path.dirname(__file__), "fonts", "THSarabunNew Bold.ttf")
        big_font_size = font_size * scale
        font = ImageFont.truetype(font_path, big_font_size)
        padding = 4 * scale
        lines = text.split('\n')

        # ใช้ fixed line height แทน bbox เพื่อหลีกเลี่ยงปัญหา tone marks ทำให้ความสูงไม่เท่ากัน
        fixed_line_height = int(font_size * 1.2 * scale)

        # วัดความกว้างเท่านั้น
        dummy_img = Image.new("RGBA", (10, 10), (255, 255, 255, 0))
        dummy_draw = ImageDraw.Draw(dummy_img)
        line_widths = []
        for line in lines:
            bbox = dummy_draw.textbbox((0, 0), line, font=font)
            width = bbox[2] - bbox[0]
            line_widths.append(width)

        max_width = max(line_widths) + 2 * padding
        total_height = len(lines) * fixed_line_height + 2 * padding

        img = Image.new("RGBA", (max_width, total_height), (255, 255, 255, 0))
        draw = ImageDraw.Draw(img)

        y = padding
        for line in lines:
            # ใช้ anchor="la" (left-ascender) เพื่อยึดตำแหน่งที่ ascender line
            # ทำให้ทุกบรรทัดวางที่ตำแหน่งเดียวกัน ไม่ว่าจะมี tone marks หรือไม่
            draw.text((padding, y), line, font=font, fill=color, anchor="la")
            y += fixed_line_height

        return img

    try:
        # ===== ส่วนที่ 1: เพิ่มลายเซ็น (จาก /add_signature_v2) =====
        DEFAULT_SIGNATURE_HEIGHT = 50
        DEFAULT_COMMENT_FONT_SIZE = 18
        font_path = os.path.join(os.path.dirname(__file__), "fonts", "THSarabunNew.ttf")
        bold_font_path = os.path.join(os.path.dirname(__file__), "fonts", "THSarabunNew Bold.ttf")

        if not os.path.isfile(font_path):
            return jsonify({'error': f"Font file not found: {font_path}"}), 500

        if 'pdf' not in request.files:
            return jsonify({'error': 'No PDF file uploaded'}), 400
        pdf_file = request.files['pdf']

        if 'signatures' not in request.form:
            return jsonify({'error': 'No signatures data'}), 400
        signatures = json.loads(request.form['signatures'])

        pdf_bytes = pdf_file.read()
        pdf = fitz.open(stream=pdf_bytes, filetype="pdf")

        from collections import defaultdict
        sig_dict = defaultdict(list)
        for sig in signatures:
            page_number = int(sig.get('page', 0))
            x = int(sig['x'])
            y = int(sig['y'])
            # รองรับ width/height สำหรับ center positioning
            width = sig.get('width', 0)
            height = sig.get('height', 0)

            # ถ้าไม่มี width/height ให้ใช้ค่า default สำหรับ center positioning
            if width == 0 and height == 0:
                width = 120  # default width
                height = 60  # default height
                print(f"DEBUG: Using default dimensions {width}x{height} for signature at ({x}, {y})")

            sig_dict[(page_number, x, y, width, height)].append(sig)

        for (page_number, x, y, width, height), sigs in sig_dict.items():
            page = pdf[page_number]

            # Debug: แสดงข้อมูล page และพิกัด
            page_rect = page.rect
            print(f"DEBUG: Page {page_number} - Size: {page_rect.width}x{page_rect.height}")
            print(f"DEBUG: Original coordinates: ({x}, {y})")
            print(f"DEBUG: Signature dimensions: {width}x{height}")
            print(f"DEBUG: Page bounds: x(0-{page_rect.width}), y(0-{page_rect.height})")

            # ถ้ามี width/height แสดงว่าเป็น center positioning
            is_center_positioning = width > 0 and height > 0

            # Logic สำหรับปรับพิกัด Y ให้กลับกัน (บนเป็นล่าง ล่างเป็นบน)
            if is_center_positioning:
                # สำหรับ center positioning ใช้ height ที่กำหนดมา
                adjusted_y = page_rect.height - y - height
                # เลื่อนลงแนวดิ่งเท่ากับ height (60)
                adjusted_y += height+30
                print(f"DEBUG: Y-axis flip with center positioning: {y} -> {adjusted_y} (with +{height} offset)")
                center_x = x
                center_y = adjusted_y
                print(f"DEBUG: Using center positioning - adjusted coordinates")
                print(f"DEBUG: Center point: ({center_x}, {center_y})")
                print(f"DEBUG: Bounding box: {width}x{height}")
            else:
                # สำหรับ top-left positioning ใช้ default signature height
                signature_box_height = 60  # default height สำหรับการคำนวณ
                adjusted_y = page_rect.height - y - signature_box_height
                # เลื่อนลงแนวดิ่งเท่ากับ 60
                adjusted_y += 60
                print(f"DEBUG: Y-axis flip with top-left positioning: {y} -> {adjusted_y} (with +60 offset)")
                center_x = x
                center_y = adjusted_y
                print(f"DEBUG: Using top-left positioning - adjusted coordinates")

            current_y = center_y  # ใช้ค่า Y ที่ปรับแล้ว
            # Check if any signature has 'lines' field
            has_lines = any('lines' in sig for sig in sigs)
            if has_lines:
                # For each sig with lines, draw lines in order
                for sig in sigs:
                    lines = sig.get('lines')
                    if not lines:
                        # fallback to old logic for this sig
                        if sig['type'] == 'text':
                            text = to_thai_digits(sig.get('text', ''))
                            font_size = 20 if sig.get('type') == 'comment' else DEFAULT_COMMENT_FONT_SIZE
                            font_weight = "bold" if sig.get('type') == 'comment' else "regular"
                            orig_color = sig.get('color', (2, 53, 139))
                            if isinstance(orig_color, (list, tuple)):
                                r = min(int(orig_color[0]*0.8), 255)
                                g = min(int(orig_color[1]*0.8), 255)
                                b = min(int(orig_color[2]*0.8), 255)
                                color = (r, g, b)
                            else:
                                color = (2, 53, 139)
                            img = draw_text_image_v2(text, font_path, font_size=font_size, color=color, scale=1, font_weight=font_weight)
                            img_byte_arr = io.BytesIO()
                            img.save(img_byte_arr, format='PNG')
                            # ใช้ center positioning ถ้ามี width/height
                            if is_center_positioning:
                                left_x = center_x - img.width // 2
                                top_y = center_y - img.height // 2
                            else:
                                left_x = x
                                top_y = current_y
                            rect = fitz.Rect(left_x, top_y, left_x + img.width, top_y + img.height)
                            print(f"DEBUG: Text '{text}' placed at rect: {rect} (center_pos: {is_center_positioning})")
                            page.insert_image(rect, stream=img_byte_arr.getvalue())
                            if not is_center_positioning:
                                current_y += img.height
                        elif sig['type'] == 'image':
                            file_key = sig['file_key']
                            if file_key not in request.files:
                                continue
                            signature_file = request.files[file_key]
                            img = Image.open(signature_file)
                            fixed_height = DEFAULT_SIGNATURE_HEIGHT
                            ratio = fixed_height / img.height
                            new_width = int(img.width * ratio)
                            img = img.resize((new_width, fixed_height), resample=Image.LANCZOS)
                            img_byte_arr = io.BytesIO()
                            img.save(img_byte_arr, format='PNG')
                            # ใช้ center positioning ถ้ามี width/height
                            if is_center_positioning:
                                left_x = center_x - new_width // 2
                                top_y = center_y - fixed_height // 2
                            else:
                                left_x = x
                                top_y = current_y
                            rect = fitz.Rect(left_x, top_y, left_x + new_width, top_y + fixed_height)
                            print(f"DEBUG: Image placed at rect: {rect} (center_pos: {is_center_positioning})")
                            page.insert_image(rect, stream=img_byte_arr.getvalue())
                            if not is_center_positioning:
                                current_y += fixed_height
                    else:
                        # draw lines in order - รองรับ center positioning
                        if is_center_positioning:
                            # สำหรับ center positioning ให้เริ่มจากด้านบนของ bounding box
                            current_y = center_y - height // 2

                        for line in lines:
                            line_type = line.get('type')
                            if line_type == 'image':
                                file_key = line.get('file_key')
                                if file_key and file_key in request.files:
                                    signature_file = request.files[file_key]
                                    img = Image.open(signature_file)
                                    fixed_height = DEFAULT_SIGNATURE_HEIGHT
                                    ratio = fixed_height / img.height
                                    new_width = int(img.width * ratio)
                                    img = img.resize((new_width, fixed_height), resample=Image.LANCZOS)
                                    img_byte_arr = io.BytesIO()
                                    img.save(img_byte_arr, format='PNG')

                                    if is_center_positioning:
                                        left_x = center_x - new_width // 2
                                        top_y = current_y
                                        print(f"DEBUG: Image center positioning - center_x:{center_x}, new_width:{new_width}, left_x:{left_x}")
                                        print(f"DEBUG: Expected position - should place image at left edge: {left_x}")
                                    else:
                                        left_x = x
                                        top_y = current_y

                                    rect = fitz.Rect(left_x, top_y, left_x + new_width, top_y + fixed_height)
                                    print(f"DEBUG: Image rect: {rect}")
                                    page.insert_image(rect, stream=img_byte_arr.getvalue())
                                    current_y += fixed_height
                            else:
                                # For text types: 'comment', 'name', 'position', 'academic_rank', 'org_structure_role', 'timestamp'
                                text_value = line.get('text') or line.get('value') or line.get('comment') or ''

                                # ฟังก์ชันนับตัวอักษรที่มองเห็น (ไม่รวม tone marks, vowel marks)
                                def count_visible_chars(s):
                                    thai_marks = set([
                                        '\u0E31', '\u0E34', '\u0E35', '\u0E36', '\u0E37',
                                        '\u0E38', '\u0E39', '\u0E3A', '\u0E47', '\u0E48',
                                        '\u0E49', '\u0E4A', '\u0E4B', '\u0E4C', '\u0E4D', '\u0E4E'
                                    ])
                                    return len([c for c in s if c not in thai_marks])

                                # ฟังก์ชันตัดข้อความตามจำนวนตัวอักษรที่มองเห็น
                                def wrap_by_visible_chars(text, max_chars=30):
                                    if count_visible_chars(text) <= max_chars:
                                        return [text]

                                    lines = []
                                    current = ""
                                    for char in text:
                                        test = current + char
                                        if count_visible_chars(test) <= max_chars:
                                            current = test
                                        else:
                                            if current:
                                                lines.append(current)
                                            current = char
                                    if current:
                                        lines.append(current)
                                    return lines

                                # ถ้าเป็น comment และมี - ให้แยกเป็นหลายบรรทัด
                                if line_type == 'comment' and '-' in text_value:
                                    # แยกตาม - และเก็บ - ไว้ข้างหน้าแต่ละบรรทัด
                                    parts = text_value.split('-')
                                    # เอา parts ที่ไม่ว่างเปล่าออกมา
                                    text_lines = ['-' + part for part in parts if part.strip()]

                                    # ตัดแต่ละบรรทัดให้ไม่เกิน 30 ตัวอักษรที่มองเห็น
                                    wrapped_lines = []
                                    for tline in text_lines:
                                        wrapped = wrap_by_visible_chars(tline, max_chars=30)
                                        wrapped_lines.extend(wrapped)

                                    text = '\n'.join([to_thai_digits(t) for t in wrapped_lines])
                                else:
                                    text = to_thai_digits(text_value)
                                font_size = 20 if line_type == 'comment' else DEFAULT_COMMENT_FONT_SIZE
                                font_weight = "bold" if line_type == 'comment' else "regular"
                                orig_color = line.get('color', (2, 53, 139))
                                if isinstance(orig_color, (list, tuple)):
                                    r = min(int(orig_color[0]*0.8), 255)
                                    g = min(int(orig_color[1]*0.8), 255)
                                    b = min(int(orig_color[2]*0.8), 255)
                                    color = (r, g, b)
                                else:
                                    color = (2, 53, 139)
                                img = draw_text_image_v2(text, font_path, font_size=font_size, color=color, scale=1, font_weight=font_weight)
                                img_byte_arr = io.BytesIO()
                                img.save(img_byte_arr, format='PNG')

                                if is_center_positioning:
                                    left_x = center_x - img.width // 2
                                    top_y = current_y
                                    print(f"DEBUG: Text center positioning - center_x:{center_x}, img.width:{img.width}, left_x:{left_x}")
                                    print(f"DEBUG: Expected position - should place text at left edge: {left_x}")
                                else:
                                    left_x = x
                                    top_y = current_y

                                rect = fitz.Rect(left_x, top_y, left_x + img.width, top_y + img.height)
                                print(f"DEBUG: Text '{text}' rect: {rect}")
                                page.insert_image(rect, stream=img_byte_arr.getvalue())
                                current_y += img.height
            else:
                # fallback to old logic
                sigs_sorted = sorted(sigs, key=lambda s: 0 if s['type'] == 'text' else 1)
                for sig in sigs_sorted:
                    if sig['type'] == 'text':
                        text = to_thai_digits(sig.get('text', ''))
                        font_size = 20 if sig.get('type') == 'comment' else DEFAULT_COMMENT_FONT_SIZE
                        font_weight = "bold" if sig.get('type') == 'comment' else "regular"
                        orig_color = sig.get('color', (2, 53, 139))
                        if isinstance(orig_color, (list, tuple)):
                            r = min(int(orig_color[0]*0.8), 255)
                            g = min(int(orig_color[1]*0.8), 255)
                            b = min(int(orig_color[2]*0.8), 255)
                            color = (r, g, b)
                        else:
                            color = (2, 53, 139)
                        img = draw_text_image_v2(text, font_path, font_size=font_size, color=color, scale=1, font_weight=font_weight)
                        img_byte_arr = io.BytesIO()
                        img.save(img_byte_arr, format='PNG')
                        rect = fitz.Rect(x, current_y, x + img.width, current_y + img.height)
                        page.insert_image(rect, stream=img_byte_arr.getvalue())
                        current_y += img.height
                    elif sig['type'] == 'image':
                        file_key = sig['file_key']
                        if file_key not in request.files:
                            continue
                        signature_file = request.files[file_key]
                        img = Image.open(signature_file)
                        fixed_height = DEFAULT_SIGNATURE_HEIGHT
                        ratio = fixed_height / img.height
                        new_width = int(img.width * ratio)
                        img = img.resize((new_width, fixed_height), resample=Image.LANCZOS)
                        img_byte_arr = io.BytesIO()
                        img.save(img_byte_arr, format='PNG')
                        rect = fitz.Rect(x, current_y, x + new_width, current_y + fixed_height)
                        page.insert_image(rect, stream=img_byte_arr.getvalue())
                        current_y += fixed_height

        # ===== ส่วนที่ 2: เพิ่มตราสรุป (จาก /stamp_summary) =====
        if 'summary_payload' in request.form and 'sign_png' in request.files:
            print("[DEBUG] Adding stamp summary...")

            sign_file = request.files['sign_png']
            p = json.loads(request.form['summary_payload'])
            summary = p.get('summary', '')
            group_name = p.get('group_name', '')
            receiver_name = p.get('receiver_name', '')
            date = p.get('date', '')

            print(f"[DEBUG] Summary data: summary={summary}, group={group_name}, receiver={receiver_name}, date={date}")

            page = pdf[0]  # ใช้หน้าแรก

            # เตรียมข้อมูลสำหรับคำนวณความสูง
            page_w = page.rect.width
            page_h = page.rect.height
            margin = 30
            stamp_width = 200

            # ฟังก์ชันตัดข้อความ
            def wrap_text(text, max_chars_approx):
                thai_marks = set([
                    '\u0E31', '\u0E34', '\u0E35', '\u0E36', '\u0E37', '\u0E38', '\u0E39', '\u0E3A',
                    '\u0E47', '\u0E48', '\u0E49', '\u0E4A', '\u0E4B', '\u0E4C', '\u0E4D', '\u0E4E'
                ])

                def count_visible_chars(s):
                    return len([c for c in s if c not in thai_marks])

                def cut_at_visible_chars(s, max_visible):
                    if count_visible_chars(s) <= max_visible:
                        return s

                    visible_count = 0
                    for i, char in enumerate(s):
                        if char not in thai_marks:
                            visible_count += 1
                            if visible_count >= max_visible:
                                current_pos = i + 1
                                for j in range(current_pos, len(s)):
                                    if s[j] == ' ' or s[j] in '.!?,:;':
                                        return s[:j]
                                cut_pos = current_pos
                                while cut_pos > 0 and s[cut_pos-1] in thai_marks:
                                    cut_pos -= 1
                                return s[:max(cut_pos, current_pos-5)]
                    return s

                lines = []
                max_chars = 30

                # ถ้าขึ้นต้นด้วย "เรื่อง "
                if text.startswith("เรื่อง "):
                    words = text.split(' ')
                    if len(words) >= 2:
                        first_group = f"เรื่อง {words[1]}"
                        rest_text = ' '.join(words[2:])

                        if count_visible_chars(first_group) > max_chars:
                            text = text
                        else:
                            lines.append(first_group)
                            text = rest_text

                while count_visible_chars(text) > max_chars:
                    cut_text = cut_at_visible_chars(text, max_chars)
                    lines.append(cut_text)
                    text = text[len(cut_text):]

                if text.strip():
                    lines.append(text)

                return lines

            # ฟังก์ชันวาดข้อความสำหรับตรา (รองรับการ wrap text)
            def draw_text_img(text, size=16, bold=False, max_width=None):
                from PIL import ImageFont, ImageDraw, Image
                fp = bold_font_path if bold else font_path
                color_rgb = (2, 53, 139)
                text = to_thai_digits(text)

                # ถ้าไม่มี max_width ใช้วิธีเดิม
                if max_width is None:
                    img = draw_text_image(text, fp, size, color_rgb, scale=1)
                    return img

                # มี max_width ให้ wrap text
                font = ImageFont.truetype(fp, size)
                padding = 4

                # แยกข้อความเป็นบรรทัดตาม max_width
                words = text.split(' ')
                lines = []
                current_line = ""

                for i, word in enumerate(words):
                    test_line = current_line + (" " if current_line else "") + word
                    bbox = font.getbbox(test_line)
                    text_width = bbox[2] - bbox[0]

                    if text_width <= max_width - 2 * padding:
                        current_line = test_line
                    else:
                        # กรณีพิเศษ: ถ้าบรรทัดปัจจุบันคือ "เรื่อง" หรือ "เห็นควรมอบ" ต้องเก็บคำถัดไปไว้ด้วยเสมอ
                        if current_line == "เรื่อง" or current_line == "เห็นควรมอบ":
                            # บังคับรวม prefix กับคำแรกของข้อความ แล้วตัดทีละตัวอักษร
                            # เพิ่ม space ครั้งเดียวหลัง prefix
                            current_line = current_line + " "
                            for char in word:
                                test_char = current_line + char
                                char_bbox = font.getbbox(test_char)
                                char_width = char_bbox[2] - char_bbox[0]

                                if char_width <= max_width - 2 * padding:
                                    current_line = test_char
                                else:
                                    if current_line:
                                        lines.append(current_line)
                                    current_line = char
                        else:
                            # ถ้าคำยาวเกินไป ต้องตัดทีละตัวอักษร
                            word_bbox = font.getbbox(word)
                            word_width = word_bbox[2] - word_bbox[0]

                            if word_width > max_width - 2 * padding:
                                # คำยาวเกิน ต้องตัดทีละตัวอักษร
                                if current_line:
                                    lines.append(current_line)
                                    current_line = ""

                                # ตัดคำยาวทีละตัวอักษร
                                for char in word:
                                    test_char = current_line + char
                                    char_bbox = font.getbbox(test_char)
                                    char_width = char_bbox[2] - char_bbox[0]

                                    if char_width <= max_width - 2 * padding:
                                        current_line += char
                                    else:
                                        if current_line:
                                            lines.append(current_line)
                                        current_line = char
                            else:
                                # คำไม่ยาวเกิน แต่รวมกับบรรทัดปัจจุบันแล้วยาวเกิน
                                if current_line:
                                    lines.append(current_line)
                                current_line = word

                if current_line:
                    lines.append(current_line)

                # คำนวณขนาดภาพ
                max_line_width = 0
                fixed_line_height = 14  # ใช้ความสูงคงที่

                for line in lines:
                    bbox = font.getbbox(line)
                    line_width = bbox[2] - bbox[0]
                    max_line_width = max(max_line_width, line_width)

                img_width = max_line_width + 2 * padding
                # ใช้ความสูงคงที่ 14px ต่อบรรทัด
                img_height = len(lines) * fixed_line_height + 2 * padding

                # สร้างภาพ
                img = Image.new("RGBA", (img_width, img_height), (255, 255, 255, 0))
                draw = ImageDraw.Draw(img)

                y = padding
                for i, line in enumerate(lines):
                    draw.text((padding, y), line, font=font, fill=color_rgb)
                    # ใช้ความสูงคงที่
                    y += fixed_line_height

                # เก็บจำนวนบรรทัดไว้ใน attribute ของภาพ
                img.line_count = len(lines)
                return img

            # สร้างข้อความทั้งหมดก่อนเพื่อคำนวณความสูงจริง
            font_size = 16
            text_max_width = stamp_width - 20

            # สร้างภาพข้อความทั้งหมดก่อน
            text1 = "เรียน ผอ. ศกศ.เขต ๖ จ.ลพบุรี"
            img1 = draw_text_img(text1, size=font_size, bold=True, max_width=text_max_width)

            # ฟังก์ชันสร้างภาพข้อความแบบ mixed (prefix ตัวหนา, content ตัวปกติ)
            def draw_mixed_text_img(prefix, content, size=16, max_width=None):
                from PIL import ImageFont, ImageDraw, Image
                color_rgb = (2, 53, 139)
                padding = 4
                fixed_line_height = 14

                # แปลงเป็นเลขไทย
                prefix = to_thai_digits(prefix)
                content = to_thai_digits(content)

                bold_font = ImageFont.truetype(bold_font_path, size)
                normal_font = ImageFont.truetype(font_path, size)

                # วาด prefix (ตัวหนา) + content (ตัวปกติ) โดย wrap ทั้งหมด
                full_text = prefix + " " + content

                # แยกข้อความเป็นบรรทัดตาม max_width
                words = full_text.split(' ')
                lines = []
                current_line = ""

                for i, word in enumerate(words):
                    test_font = bold_font if i == 0 else normal_font
                    test_line = current_line + (" " if current_line else "") + word
                    bbox = test_font.getbbox(test_line) if i == 0 else normal_font.getbbox(test_line)
                    text_width = bbox[2] - bbox[0]

                    if text_width <= max_width - 2 * padding:
                        current_line = test_line
                    else:
                        if current_line == prefix:
                            current_line = current_line + " "
                            for char in word:
                                test_char = current_line + char
                                char_bbox = normal_font.getbbox(test_char)
                                char_width = char_bbox[2] - char_bbox[0]

                                if char_width <= max_width - 2 * padding:
                                    current_line = test_char
                                else:
                                    if current_line:
                                        lines.append(current_line)
                                    current_line = char
                        else:
                            word_bbox = normal_font.getbbox(word)
                            word_width = word_bbox[2] - word_bbox[0]

                            if word_width > max_width - 2 * padding:
                                if current_line:
                                    lines.append(current_line)
                                    current_line = ""

                                for char in word:
                                    test_char = current_line + char
                                    char_bbox = normal_font.getbbox(test_char)
                                    char_width = char_bbox[2] - char_bbox[0]

                                    if char_width <= max_width - 2 * padding:
                                        current_line += char
                                    else:
                                        if current_line:
                                            lines.append(current_line)
                                        current_line = char
                            else:
                                if current_line:
                                    lines.append(current_line)
                                current_line = word

                if current_line:
                    lines.append(current_line)

                # คำนวณความกว้างสูงสุด
                max_line_width = 0
                for line in lines:
                    if line.startswith(prefix):
                        prefix_bbox = bold_font.getbbox(prefix)
                        prefix_width = prefix_bbox[2] - prefix_bbox[0]
                        rest = line[len(prefix):]
                        rest_bbox = normal_font.getbbox(rest)
                        rest_width = rest_bbox[2] - rest_bbox[0]
                        line_width = prefix_width + rest_width
                    else:
                        bbox = normal_font.getbbox(line)
                        line_width = bbox[2] - bbox[0]
                    max_line_width = max(max_line_width, line_width)

                img_width = max_line_width + 2 * padding
                img_height = len(lines) * fixed_line_height + 2 * padding

                # สร้างภาพและวาดข้อความ
                img = Image.new("RGBA", (img_width, img_height), (255, 255, 255, 0))
                draw = ImageDraw.Draw(img)

                y = padding
                for line in lines:
                    if line.startswith(prefix):
                        draw.text((padding, y), prefix, font=bold_font, fill=color_rgb)
                        prefix_bbox = bold_font.getbbox(prefix)
                        prefix_width = prefix_bbox[2] - prefix_bbox[0]
                        rest = line[len(prefix):]
                        draw.text((padding + prefix_width, y), rest, font=normal_font, fill=color_rgb)
                    else:
                        draw.text((padding, y), line, font=normal_font, fill=color_rgb)
                    y += fixed_line_height

                img.line_count = len(lines)
                return img

            img_subject = draw_mixed_text_img("เรื่อง", summary, size=font_size, max_width=text_max_width)
            img_assign = draw_mixed_text_img("เห็นควรมอบ", group_name, size=font_size, max_width=text_max_width)

            sign_img_temp = Image.open(sign_file)
            sign_height = 30
            ratio = sign_height / sign_img_temp.height
            sign_width = int(sign_img_temp.width * ratio)
            sign_img = sign_img_temp.resize((sign_width, sign_height), Image.LANCZOS)

            sign_text = "ลงชื่อ"
            img_sign_text = draw_text_img(sign_text, size=font_size, bold=False)

            receiver_text_str = f"ผู้รับ  {receiver_name}"
            img_receiver = draw_text_img(receiver_text_str, size=font_size, bold=False)

            date_text_str = f"วันที่ {date}"
            img_date = draw_text_img(date_text_str, size=font_size, bold=False)

            # คำนวณความสูงจริงตามที่จะใช้ในการวาด
            padding_top = 8
            padding_bottom = 8
            line_height = 14

            total_height = padding_top
            total_height += line_height  # เรียน ผอ.
            total_height += img_subject.height + 2  # เรื่อง (ใช้ความสูงจริง)
            total_height += img_assign.height + 2  # เห็นควรมอบ (ใช้ความสูงจริง)
            total_height += line_height + 2  # ลงชื่อ
            total_height += line_height  # ผู้รับ
            total_height += line_height  # วันที่
            total_height += padding_bottom

            stamp_height = int(total_height)

            # คำนวณตำแหน่งกรอบ (มุมซ้ายล่าง)
            center_x = margin + stamp_width//2
            center_y = page_h - margin - stamp_height//2

            # วาดกรอบตรา
            box_left = center_x - stamp_width//2
            box_top = center_y - stamp_height//2
            box_right = center_x + stamp_width//2
            box_bottom = center_y + stamp_height//2

            box_rect = fitz.Rect(box_left, box_top, box_right, box_bottom)
            box_color = (2/255, 53/255, 139/255)
            page.draw_rect(box_rect, color=box_color, width=2)

            def paste_at_position(img, x, y):
                rect = fitz.Rect(x, y, x+img.width, y+img.height)
                bio = io.BytesIO()
                img.save(bio, format='PNG')
                page.insert_image(rect, stream=bio.getvalue())

            # วาดข้อความในตรา (ใช้ภาพที่สร้างไว้แล้ว)
            # ใช้ระยะห่างแบบเดิม
            line_height = 14
            current_y = box_top + padding_top

            # เรียน ผอ.
            paste_at_position(img1, box_left + 10, current_y)
            current_y += line_height

            # เรื่อง + summary
            paste_at_position(img_subject, box_left + 10, current_y)
            # ใช้ความสูงจริงของภาพ + ระยะห่างเล็กน้อย
            current_y += img_subject.height + 2

            # เห็นควรมอบ + group_name
            paste_at_position(img_assign, box_left + 10, current_y)
            current_y += img_assign.height + 2

            # ลายเซ็น (ใช้ภาพที่สร้างไว้แล้ว)
            center_x_frame = box_left + stamp_width//2

            total_width = img_sign_text.width + 5 + sign_width
            start_x = center_x_frame - total_width//2

            sign_y = current_y
            paste_at_position(img_sign_text, start_x, sign_y)
            paste_at_position(sign_img, start_x + img_sign_text.width + 5, sign_y)

            current_y += line_height + 2

            # ผู้รับ (ใช้ภาพที่สร้างไว้แล้ว)
            paste_at_position(img_receiver, center_x_frame - img_receiver.width//2, current_y)
            current_y += line_height

            # วันที่ (ใช้ภาพที่สร้างไว้แล้ว)
            paste_at_position(img_date, center_x_frame - img_date.width//2, current_y)

        # บันทึกและส่งไฟล์กลับ
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_pdf:
            pdf.save(tmp_pdf.name)
        pdf.close()

        print("[DEBUG] PDF saved, sending response...")
        response = send_file(tmp_pdf.name, mimetype="application/pdf", as_attachment=True, download_name="signed_receive.pdf")
        response.headers['X-Debug'] = 'add_signature_receive_processed'
        return response

    except Exception as e:
        print(traceback.format_exc())
        return jsonify({'error': str(e)}), 500


if __name__ == "__main__":
    # สำหรับ Railway ต้องฟังที่ 0.0.0.0
    app.run(debug=True, host="0.0.0.0", port=5000)